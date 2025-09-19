"""
LegalSaathi Document Advisor - Refactored Main Application
Clean, modular, and maintainable Flask application for legal document analysis
"""

from flask import Flask, render_template, request, jsonify, send_file, make_response
from dataclasses import dataclass
from typing import List, Dict, Any, Optional, Tuple
import time
import json
import re
import logging
import hashlib
import os
from datetime import datetime, timedelta
from functools import wraps
from io import BytesIO
import threading
from concurrent.futures import ThreadPoolExecutor
from openai import OpenAI

# Import core services
from risk_classifier import RiskClassifier, classify_document_risk, RiskLevel as ClassifierRiskLevel, UserExpertiseDetector, EnhancedRiskCategories
from google_translate_service import GoogleTranslateService
from google_document_ai_service import document_ai_service
from google_natural_language_service import natural_language_service
from google_speech_service import speech_service
from file_processor import FileProcessor
from document_classifier import DocumentClassifier, DocumentType

# Import for PDF/Word export functionality
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# Application Configuration
class AppConfig:
    """Centralized application configuration"""
    
    def __init__(self):
        self.MAX_CONTENT_LENGTH = 10 * 1024 * 1024  # 10MB
        self.CACHE_EXPIRY = 3600  # 1 hour
        self.RATE_LIMIT_REQUESTS = 10  # requests per minute
        self.RATE_LIMIT_WINDOW = 60  # seconds
        self.MAX_DOCUMENT_SIZE = 50000  # characters
        self.MIN_DOCUMENT_SIZE = 100  # characters
        
        # AI Service Configuration
        self.OPENAI_BASE_URL = "http://172.25.0.211:8002/v1"
        self.OPENAI_API_KEY = "EMPTY"
        self.OPENAI_TIMEOUT = 5.0
        
        # Logging Configuration
        self.LOG_LEVEL = logging.INFO
        self.LOG_FORMAT = '%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        self.LOG_FILE = 'legal_saathi.log'

# Initialize Flask app
app = Flask(__name__)
config = AppConfig()
app.config['MAX_CONTENT_LENGTH'] = config.MAX_CONTENT_LENGTH

# Configure logging
logging.basicConfig(
    level=config.LOG_LEVEL,
    format=config.LOG_FORMAT,
    handlers=[
        logging.FileHandler(config.LOG_FILE),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Data Models
@dataclass
class RiskLevel:
    level: str  # "RED", "YELLOW", "GREEN"
    score: float  # 0.0 to 1.0
    reasons: List[str]
    severity: str
    confidence_percentage: int  # 0-100%
    risk_categories: dict  # financial, legal, operational
    low_confidence_warning: bool

@dataclass
class AnalysisResult:
    clause_id: str
    risk_level: RiskLevel
    plain_explanation: str
    legal_implications: List[str]
    recommendations: List[str]

@dataclass
class DocumentAnalysis:
    document_id: str
    analysis_results: List[AnalysisResult]
    overall_risk: RiskLevel
    summary: str
    processing_time: float

# Service Manager
class ServiceManager:
    """Centralized service management and initialization"""
    
    def __init__(self):
        self.services = {}
        self.initialize_services()
    
    def initialize_services(self):
        """Initialize all application services"""
        try:
            # Core services
            self.services['file_processor'] = FileProcessor()
            self.services['document_classifier'] = DocumentClassifier()
            self.services['translation'] = GoogleTranslateService()
            self.services['expertise_detector'] = UserExpertiseDetector()
            self.services['risk_categories'] = EnhancedRiskCategories()
            
            # Google Cloud AI services
            self.services['google_document_ai'] = document_ai_service
            self.services['google_natural_language'] = natural_language_service
            self.services['google_speech'] = speech_service
            
            # OpenAI client for vLLM server
            self.services['openai_client'] = OpenAI(
                base_url=config.OPENAI_BASE_URL,
                api_key=config.OPENAI_API_KEY,
                timeout=config.OPENAI_TIMEOUT
            )
            
            # AI clarification service
            self.services['ai_clarification'] = RealAIService(self.services['openai_client'])
            
            # Metrics and caching
            self.services['metrics_tracker'] = AnalysisMetrics()
            self.services['cache_manager'] = CacheManager()
            self.services['rate_limiter'] = RateLimiter()
            
            # Report exporter
            self.services['report_exporter'] = ReportExporter()
            
            logger.info("All services initialized successfully")
            
        except Exception as e:
            logger.error(f"Service initialization error: {e}")
            raise

    def get_service(self, service_name: str):
        """Get a service by name"""
        return self.services.get(service_name)

# Cache Management
class CacheManager:
    """Improved cache management with better organization"""
    
    def __init__(self):
        self.analysis_cache = {}
        self.cache_expiry = config.CACHE_EXPIRY
    
    def get_cache_key(self, text: str, analysis_type: str = "analysis") -> str:
        """Generate cache key for analysis results"""
        text_hash = hashlib.md5(text.encode('utf-8')).hexdigest()
        return f"{analysis_type}:{text_hash}"
    
    def get_from_cache(self, cache_key: str) -> Optional[Dict]:
        """Get analysis result from cache if not expired"""
        if cache_key in self.analysis_cache:
            cached_data = self.analysis_cache[cache_key]
            if time.time() - cached_data['timestamp'] < self.cache_expiry:
                logger.info(f"Cache hit for key: {cache_key}")
                return cached_data['data']
            else:
                del self.analysis_cache[cache_key]
                logger.info(f"Cache expired for key: {cache_key}")
        return None
    
    def store_in_cache(self, cache_key: str, data: Dict) -> None:
        """Store analysis result in cache"""
        self.analysis_cache[cache_key] = {
            'data': data,
            'timestamp': time.time()
        }
        logger.info(f"Stored in cache: {cache_key}")
    
    def clear_expired_cache(self):
        """Clear expired cache entries"""
        current_time = time.time()
        expired_keys = [
            key for key, value in self.analysis_cache.items()
            if current_time - value['timestamp'] > self.cache_expiry
        ]
        for key in expired_keys:
            del self.analysis_cache[key]
        
        if expired_keys:
            logger.info(f"Cleared {len(expired_keys)} expired cache entries")

# Rate Limiting
class RateLimiter:
    """Improved rate limiting with better tracking"""
    
    def __init__(self):
        self.rate_limit_cache = {}
        self.requests_per_window = config.RATE_LIMIT_REQUESTS
        self.window_size = config.RATE_LIMIT_WINDOW
    
    def is_rate_limited(self, client_ip: str) -> bool:
        """Check if client has exceeded rate limit"""
        current_time = time.time()
        
        # Clean old entries
        self._cleanup_old_entries(current_time)
        
        # Check current client
        if client_ip not in self.rate_limit_cache:
            self.rate_limit_cache[client_ip] = []
        
        client_requests = self.rate_limit_cache[client_ip]
        
        if len(client_requests) >= self.requests_per_window:
            logger.warning(f"Rate limit exceeded for IP: {client_ip}")
            return True
        
        client_requests.append(current_time)
        return False
    
    def _cleanup_old_entries(self, current_time: float):
        """Remove old rate limit entries"""
        for ip in list(self.rate_limit_cache.keys()):
            self.rate_limit_cache[ip] = [
                timestamp for timestamp in self.rate_limit_cache[ip]
                if current_time - timestamp < self.window_size
            ]
            if not self.rate_limit_cache[ip]:
                del self.rate_limit_cache[ip]

# Metrics Tracking
class AnalysisMetrics:
    """Enhanced metrics tracking with thread safety"""
    
    def __init__(self):
        self.metrics = {
            'total_analyses': 0,
            'successful_analyses': 0,
            'failed_analyses': 0,
            'cache_hits': 0,
            'avg_processing_time': 0.0,
            'confidence_scores': [],
            'error_types': {}
        }
        self.lock = threading.Lock()
    
    def record_analysis(self, success: bool, processing_time: float, confidence: float = None, error_type: str = None):
        """Record analysis metrics thread-safely"""
        with self.lock:
            self.metrics['total_analyses'] += 1
            if success:
                self.metrics['successful_analyses'] += 1
                if confidence is not None:
                    self.metrics['confidence_scores'].append(confidence)
            else:
                self.metrics['failed_analyses'] += 1
                if error_type:
                    self.metrics['error_types'][error_type] = self.metrics['error_types'].get(error_type, 0) + 1
            
            # Update average processing time
            total_time = self.metrics['avg_processing_time'] * (self.metrics['total_analyses'] - 1) + processing_time
            self.metrics['avg_processing_time'] = total_time / self.metrics['total_analyses']
    
    def get_metrics(self) -> Dict:
        """Get current metrics thread-safely"""
        with self.lock:
            metrics_copy = self.metrics.copy()
            if metrics_copy['confidence_scores']:
                metrics_copy['avg_confidence'] = sum(metrics_copy['confidence_scores']) / len(metrics_copy['confidence_scores'])
            else:
                metrics_copy['avg_confidence'] = 0.0
            return metrics_copy

# Document Validator
class DocumentValidator:
    """Centralized document validation logic"""
    
    @staticmethod
    def validate_document_input(document_text: str) -> Tuple[bool, Optional[str]]:
        """Validate document input for universal legal document analysis"""
        if not document_text or not document_text.strip():
            return False, "Please provide a valid legal document text."
        
        text = document_text.strip()
        length = len(text)
        
        # Check length restrictions
        if length < config.MIN_DOCUMENT_SIZE:
            return False, f"Document text must be at least {config.MIN_DOCUMENT_SIZE} characters long for meaningful analysis."
        
        if length > config.MAX_DOCUMENT_SIZE:
            return False, f"Document text exceeds maximum length of {config.MAX_DOCUMENT_SIZE} characters. Please shorten your document."
        
        # Basic content validation - check for general legal document keywords
        legal_keywords = [
            'agreement', 'contract', 'party', 'parties', 'terms', 'conditions',
            'shall', 'hereby', 'whereas', 'therefore', 'obligations', 'rights',
            'liability', 'breach', 'termination', 'effective', 'binding', 'legal'
        ]
        
        text_lower = text.lower()
        found_keywords = [keyword for keyword in legal_keywords if keyword in text_lower]
        
        if len(found_keywords) < 2:
            return False, "This doesn't appear to be a legal document. Please ensure you've pasted the correct document text."
        
        return True, None

# AI Service Integration
class AIServiceOrchestrator:
    """Orchestrates multiple AI services with intelligent fallback"""
    
    def __init__(self, service_manager: ServiceManager):
        self.service_manager = service_manager
        self.google_document_ai = service_manager.get_service('google_document_ai')
        self.google_natural_language = service_manager.get_service('google_natural_language')
        self.openai_client = service_manager.get_service('openai_client')
    
    def analyze_document_with_ai(self, document_text: str, document_type: DocumentType, 
                                expertise_profile: Dict = None, document_bytes: bytes = None) -> DocumentAnalysis:
        """Enhanced legal document analysis using multiple AI services"""
        start_time = time.time()
        
        try:
            # Enhanced analysis with Google Cloud AI services
            enhanced_insights = self._gather_ai_insights(document_text, document_bytes)
            
            # Continue with existing risk classification
            risk_analysis = classify_document_risk(document_text, document_type)
            
            # Convert to enhanced data model format
            analysis_results = self._convert_risk_analysis(risk_analysis, document_type, expertise_profile)
            
            # Generate enhanced summary
            summary = self._generate_summary(risk_analysis, enhanced_insights)
            
            # Create document analysis object
            document_analysis = DocumentAnalysis(
                document_id=hashlib.md5(document_text.encode()).hexdigest()[:12],
                analysis_results=analysis_results,
                overall_risk=self._convert_overall_risk(risk_analysis['overall_risk']),
                summary=summary,
                processing_time=time.time() - start_time
            )
            
            return document_analysis
            
        except Exception as e:
            logger.error(f"AI analysis error: {e}")
            raise
    
    def _gather_ai_insights(self, document_text: str, document_bytes: bytes = None) -> Dict:
        """Gather insights from multiple AI services in parallel"""
        enhanced_insights = {}
        
        with ThreadPoolExecutor(max_workers=3) as executor:
            futures = {}
            
            # Google Cloud Document AI analysis (if document bytes available)
            if document_bytes and self.google_document_ai.enabled:
                futures['document_ai'] = executor.submit(
                    self.google_document_ai.process_legal_document, document_bytes
                )
            
            # Google Cloud Natural Language AI analysis
            if self.google_natural_language.enabled:
                futures['natural_language'] = executor.submit(
                    self.google_natural_language.analyze_legal_document, document_text
                )
            
            # Process results
            for service, future in futures.items():
                try:
                    result = future.result(timeout=30)
                    if result.get('success'):
                        enhanced_insights[service] = result
                        logger.info(f"{service} analysis completed successfully")
                except Exception as e:
                    logger.warning(f"{service} analysis failed: {e}")
        
        return enhanced_insights
    
    def _convert_risk_analysis(self, risk_analysis: Dict, document_type: DocumentType, 
                              expertise_profile: Dict) -> List[AnalysisResult]:
        """Convert risk analysis to enhanced format with adaptive explanations"""
        analysis_results = []
        doc_context = self._get_document_context(document_type)
        
        for clause_assessment in risk_analysis['clause_assessments']:
            assessment = clause_assessment['assessment']
            
            # Convert classifier risk level to enhanced format
            risk_level = RiskLevel(
                level=assessment.level.value,
                score=assessment.score,
                reasons=assessment.reasons,
                severity=assessment.severity,
                confidence_percentage=assessment.confidence_percentage,
                risk_categories=assessment.risk_categories,
                low_confidence_warning=assessment.low_confidence_warning
            )
            
            # Generate adaptive explanations
            explanation, implications, recommendations = self._generate_adaptive_explanation(
                assessment, doc_context, expertise_profile
            )
            
            analysis_result = AnalysisResult(
                clause_id=clause_assessment['clause_id'],
                risk_level=risk_level,
                plain_explanation=explanation,
                legal_implications=implications,
                recommendations=recommendations
            )
            analysis_results.append(analysis_result)
        
        return analysis_results
    
    def _get_document_context(self, document_type: DocumentType) -> Dict[str, str]:
        """Get context-specific terms for different document types"""
        contexts = {
            DocumentType.RENTAL_AGREEMENT: {
                'party_disadvantage': 'tenant disadvantage',
                'other_party': 'landlord',
                'protection_laws': 'tenant protection statutes',
                'legal_area': 'rental/housing'
            },
            DocumentType.EMPLOYMENT_CONTRACT: {
                'party_disadvantage': 'employee disadvantage',
                'other_party': 'employer',
                'protection_laws': 'employment protection laws',
                'legal_area': 'employment'
            },
            DocumentType.NDA: {
                'party_disadvantage': 'receiving party disadvantage',
                'other_party': 'disclosing party',
                'protection_laws': 'confidentiality regulations',
                'legal_area': 'intellectual property'
            },
            DocumentType.LOAN_AGREEMENT: {
                'party_disadvantage': 'borrower disadvantage',
                'other_party': 'lender',
                'protection_laws': 'consumer lending laws',
                'legal_area': 'financial/lending'
            },
            DocumentType.PARTNERSHIP_AGREEMENT: {
                'party_disadvantage': 'partner disadvantage',
                'other_party': 'other partners',
                'protection_laws': 'partnership regulations',
                'legal_area': 'business/corporate'
            }
        }
        
        return contexts.get(document_type, {
            'party_disadvantage': 'party disadvantage',
            'other_party': 'other party',
            'protection_laws': 'applicable regulations',
            'legal_area': 'relevant legal'
        })
    
    def _generate_adaptive_explanation(self, assessment, doc_context: Dict, expertise_profile: Dict) -> Tuple[str, List[str], List[str]]:
        """Generate explanations adapted to user expertise level"""
        confidence_indicator = ""
        if assessment.low_confidence_warning:
            confidence_indicator = f" (⚠️ Low Confidence: {assessment.confidence_percentage}%)"
        
        expertise_level = expertise_profile.get('level', 'beginner') if expertise_profile else 'beginner'
        
        if expertise_level == 'expert':
            return self._generate_expert_explanation(assessment, doc_context, confidence_indicator)
        elif expertise_level == 'intermediate':
            return self._generate_intermediate_explanation(assessment, doc_context, confidence_indicator)
        else:
            return self._generate_beginner_explanation(assessment, doc_context, confidence_indicator)
    
    def _generate_expert_explanation(self, assessment, doc_context: Dict, confidence_indicator: str) -> Tuple[str, List[str], List[str]]:
        """Generate technical explanations for experts"""
        if assessment.level == ClassifierRiskLevel.RED:
            explanation = f"🚨 HIGH RISK{confidence_indicator}: Legal analysis indicates significant {doc_context['party_disadvantage']}. Risk categories: {', '.join([k for k, v in assessment.risk_categories.items() if v > 0.6])}. {' '.join(assessment.reasons[:2])}"
            implications = [f"Potential legal liability exposure", f"Contractual imbalance favoring {doc_context['other_party']}", f"May violate {doc_context['protection_laws']}"]
            recommendations = ["Seek legal counsel for clause modification", f"Review relevant {doc_context['legal_area']} laws", "Consider contract renegotiation or withdrawal"]
        elif assessment.level == ClassifierRiskLevel.YELLOW:
            explanation = f"⚠️ MODERATE RISK{confidence_indicator}: Analysis shows suboptimal terms with negotiation potential. {' '.join(assessment.reasons[:2])}"
            implications = ["Terms deviate from market standards", f"Potential for improved {doc_context['party_disadvantage'].split()[0]} protections"]
            recommendations = ["Negotiate more balanced terms", f"Reference {doc_context['legal_area']} market standards", "Document any agreed modifications"]
        else:
            explanation = f"✅ ACCEPTABLE TERMS{confidence_indicator}: Clause aligns with standard {doc_context['legal_area']} practices. {' '.join(assessment.reasons[:1]) if assessment.reasons else 'No significant legal concerns identified.'}"
            implications = ["Terms within acceptable legal parameters", "Standard market practice"]
            recommendations = ["Terms are legally sound as written"]
        
        return explanation, implications, recommendations
    
    def _generate_intermediate_explanation(self, assessment, doc_context: Dict, confidence_indicator: str) -> Tuple[str, List[str], List[str]]:
        """Generate balanced explanations for intermediate users"""
        if assessment.level == ClassifierRiskLevel.RED:
            explanation = f"🚨 HIGH RISK{confidence_indicator}: This clause is unfavorable and should be addressed. Main concerns: {', '.join(assessment.reasons[:2])}"
            implications = ["This puts you at a significant disadvantage", "Could lead to financial or legal problems", f"Not typical in fair {doc_context['legal_area']} agreements"]
            recommendations = ["Negotiate to change this clause", "Get legal advice before signing", f"Consider this a red flag for the {doc_context['other_party']}"]
        elif assessment.level == ClassifierRiskLevel.YELLOW:
            explanation = f"⚠️ MODERATE CONCERN{confidence_indicator}: This clause could be more balanced. Issues: {' '.join(assessment.reasons[:2])}"
            implications = ["Terms are somewhat one-sided", f"May cause issues during the {doc_context['legal_area']} relationship", "Room for improvement through negotiation"]
            recommendations = ["Try to negotiate better terms", "Understand what you're agreeing to", f"Compare with other {doc_context['legal_area']} options"]
        else:
            explanation = f"✅ FAIR TERMS{confidence_indicator}: This appears to be a reasonable and standard clause. {' '.join(assessment.reasons[:1]) if assessment.reasons else 'No major concerns identified.'}"
            implications = ["Terms are balanced and fair", f"Follows standard {doc_context['legal_area']} practices"]
            recommendations = ["This clause is acceptable"]
        
        return explanation, implications, recommendations
    
    def _generate_beginner_explanation(self, assessment, doc_context: Dict, confidence_indicator: str) -> Tuple[str, List[str], List[str]]:
        """Generate simple explanations for beginners"""
        if assessment.level == ClassifierRiskLevel.RED:
            explanation = f"🚨 DANGER{confidence_indicator}: This part of the contract is bad for you! It means: {' '.join(assessment.reasons[:1])}"
            implications = ["This could cost you money or cause problems", "You might lose important rights", f"This is not fair in {doc_context['legal_area']} agreements"]
            recommendations = [f"Ask the {doc_context['other_party']} to change this", f"Get help from someone who knows about {doc_context['legal_area']}", "Don't sign if they won't change it"]
        elif assessment.level == ClassifierRiskLevel.YELLOW:
            explanation = f"⚠️ BE CAREFUL{confidence_indicator}: This part could be better for you. Problem: {' '.join(assessment.reasons[:1])}"
            implications = ["This might cause problems later", "It's not the worst, but not great either"]
            recommendations = ["Ask if this can be changed", "Make sure you understand what this means", "Think about whether you're okay with this"]
        else:
            explanation = f"✅ LOOKS GOOD{confidence_indicator}: This part seems fair and normal. {' '.join(assessment.reasons[:1]) if assessment.reasons else 'Nothing to worry about here.'}"
            implications = [f"This is typical for {doc_context['legal_area']} contracts", f"This part protects both parties involved"]
            recommendations = ["This part is fine to agree to"]
        
        return explanation, implications, recommendations
    
    def _convert_overall_risk(self, overall_risk_data: Dict) -> RiskLevel:
        """Convert overall risk data to RiskLevel object"""
        return RiskLevel(
            level=overall_risk_data['level'],
            score=overall_risk_data['score'],
            reasons=[f"Overall document risk: {overall_risk_data['severity']}"],
            severity=overall_risk_data['severity'],
            confidence_percentage=overall_risk_data['confidence_percentage'],
            risk_categories=overall_risk_data['risk_categories'],
            low_confidence_warning=overall_risk_data['low_confidence_warning']
        )
    
    def _generate_summary(self, risk_analysis: Dict, enhanced_insights: Dict) -> str:
        """Generate enhanced summary with AI insights"""
        overall_risk_data = risk_analysis['overall_risk']
        red_clauses = len([c for c in risk_analysis['clause_assessments'] if c['assessment'].level == ClassifierRiskLevel.RED])
        yellow_clauses = len([c for c in risk_analysis['clause_assessments'] if c['assessment'].level == ClassifierRiskLevel.YELLOW])
        
        confidence_note = ""
        if overall_risk_data['low_confidence_warning']:
            confidence_note = f" (⚠️ Analysis confidence: {overall_risk_data['confidence_percentage']}% - consider professional review)"
        
        # Base summary
        if overall_risk_data['level'] == 'RED':
            base_summary = f"🚨 HIGH RISK DOCUMENT{confidence_note}: This document contains {red_clauses} high-risk clauses that require immediate attention before signing."
        elif overall_risk_data['level'] == 'YELLOW':
            base_summary = f"⚠️ MODERATE RISK DOCUMENT{confidence_note}: This document has {yellow_clauses} clauses that could be improved through negotiation."
        else:
            base_summary = f"✅ LOW RISK DOCUMENT{confidence_note}: This document appears to have fair and balanced terms overall."
        
        # Add AI insights if available
        ai_insights = []
        if 'natural_language' in enhanced_insights:
            nl_data = enhanced_insights['natural_language']
            if 'legal_insights' in nl_data:
                complexity = nl_data['legal_insights'].get('document_complexity', 'unknown')
                if complexity != 'unknown':
                    ai_insights.append(f"Document complexity: {complexity}")
        
        if ai_insights:
            return f"{base_summary} AI analysis indicates: {', '.join(ai_insights)}."
        else:
            return base_summary

# Enhanced AI Clarification Service
class RealAIService:
    """Enhanced AI service for document clarification"""
    
    def __init__(self, client):
        self.client = client
        self.conversation_history = []
        self.fallback_responses = {
            'general': "I understand you're asking about your legal document. While I'm having trouble accessing my full AI capabilities right now, I'd recommend reviewing the specific clause you're concerned about and considering consultation with a legal professional if you have serious concerns.",
            'terms': "Legal terms can be complex. For specific definitions and implications, I recommend consulting with a qualified legal professional who can provide accurate interpretations based on your jurisdiction.",
            'rights': "Understanding your rights is important. While I can provide general guidance, specific rights can vary by location and situation. Consider consulting with a legal professional for personalized advice.",
            'obligations': "Legal obligations should be clearly understood before signing any document. If you're unsure about specific requirements, it's best to seek professional legal advice."
        }
    
    def ask_clarification(self, question: str, context: Dict = None) -> Dict[str, Any]:
        """Enhanced AI-powered conversational clarification"""
        try:
            # Build enhanced context from document analysis
            context_text = self._build_context_text(context)
            
            # Enhanced prompt with better structure
            prompt = self._build_clarification_prompt(question, context_text)
            
            # Make API call with enhanced error handling
            response = self.client.chat.completions.create(
                model="meta-llama/Llama-3.1-8B-Instruct",
                messages=[
                    {"role": "system", "content": "You are LegalSaathi, a helpful and knowledgeable legal document advisor focused on accessibility and user education."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=250,
                temperature=0.6,
                timeout=8.0
            )
            
            ai_response = response.choices[0].message.content.strip()
            
            # Validate response quality
            if len(ai_response) < 20:
                raise ValueError("Response too short")
            
            # Store in conversation history
            conversation_entry = self._create_conversation_entry(question, ai_response, context)
            self.conversation_history.append(conversation_entry)
            
            # Keep conversation history manageable
            if len(self.conversation_history) > 50:
                self.conversation_history = self.conversation_history[-25:]
            
            logger.info(f"AI clarification successful: {len(ai_response)} chars")
            
            return {
                'success': True,
                'response': ai_response,
                'conversation_id': len(self.conversation_history),
                'confidence_score': 85,
                'response_quality': 'high',
                'processing_time': time.time() - conversation_entry['timestamp']
            }
            
        except Exception as e:
            logger.warning(f"AI clarification error: {e}, using fallback")
            return self._get_fallback_response(question, context)
    
    def _build_context_text(self, context: Dict) -> str:
        """Build context text from document analysis"""
        if not context:
            return ""
        
        risk_level = context.get('risk_level', 'unknown')
        summary = context.get('summary', '')
        document_type = context.get('document_type', 'legal document')
        
        return f"""Document Context:
- Type: {document_type}
- Risk Level: {risk_level}
- Summary: {summary[:200]}...

"""
    
    def _build_clarification_prompt(self, question: str, context_text: str) -> str:
        """Build enhanced clarification prompt"""
        return f"""You are LegalSaathi, an AI legal document advisor specializing in making complex legal language accessible to everyday citizens.

{context_text}User Question: "{question}"

Please provide a helpful response that:
1. Directly answers the user's question in plain language
2. Explains any legal terms in simple words
3. Provides practical, actionable advice when appropriate
4. Indicates when professional legal consultation is recommended
5. Maintains a supportive, educational tone

Keep your response clear, concise (under 200 words), and focused on helping the user understand their document better."""
    
    def _create_conversation_entry(self, question: str, response: str, context: Dict) -> Dict:
        """Create conversation history entry"""
        return {
            'question': question,
            'response': response,
            'timestamp': time.time(),
            'context_provided': bool(context),
            'response_length': len(response),
            'confidence': 'high'
        }
    
    def _get_fallback_response(self, question: str, context: Dict) -> Dict:
        """Get intelligent fallback response"""
        fallback_response = self._get_intelligent_fallback(question)
        
        fallback_entry = {
            'question': question,
            'response': fallback_response,
            'timestamp': time.time(),
            'context_provided': bool(context),
            'response_length': len(fallback_response),
            'confidence': 'low',
            'fallback': True
        }
        self.conversation_history.append(fallback_entry)
        
        return {
            'success': True,
            'response': fallback_response,
            'fallback': True,
            'conversation_id': len(self.conversation_history),
            'confidence_score': 45,
            'response_quality': 'basic',
            'error_type': 'AI_SERVICE_UNAVAILABLE'
        }
    
    def _get_intelligent_fallback(self, question: str) -> str:
        """Provide intelligent fallback responses based on question content"""
        question_lower = question.lower()
        
        if any(word in question_lower for word in ['what', 'mean', 'definition', 'explain']):
            return self.fallback_responses['terms']
        elif any(word in question_lower for word in ['right', 'rights', 'can i', 'allowed']):
            return self.fallback_responses['rights']
        elif any(word in question_lower for word in ['must', 'have to', 'required', 'obligation']):
            return self.fallback_responses['obligations']
        else:
            return self.fallback_responses['general']

# Report Exporter
class ReportExporter:
    """Enhanced report export functionality"""
    
    def __init__(self):
        self.pdf_available = REPORTLAB_AVAILABLE
        self.docx_available = PYTHON_DOCX_AVAILABLE
    
    def export_analysis_to_pdf(self, analysis: DocumentAnalysis) -> BytesIO:
        """Export analysis to PDF format"""
        if not self.pdf_available:
            raise ImportError("ReportLab not available for PDF export")
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            textColor=colors.darkblue
        )
        story.append(Paragraph("LegalSaathi Document Analysis Report", title_style))
        story.append(Spacer(1, 12))
        
        # Summary
        story.append(Paragraph("Executive Summary", styles['Heading2']))
        story.append(Paragraph(analysis.summary, styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Overall Risk
        risk_color = colors.red if analysis.overall_risk.level == 'RED' else colors.orange if analysis.overall_risk.level == 'YELLOW' else colors.green
        risk_style = ParagraphStyle('RiskStyle', parent=styles['Normal'], textColor=risk_color, fontSize=12, spaceAfter=6)
        story.append(Paragraph(f"Overall Risk Level: {analysis.overall_risk.level}", risk_style))
        story.append(Spacer(1, 12))
        
        # Detailed Analysis
        story.append(Paragraph("Detailed Clause Analysis", styles['Heading2']))
        
        for i, result in enumerate(analysis.analysis_results, 1):
            # Clause header
            clause_style = ParagraphStyle('ClauseStyle', parent=styles['Heading3'], fontSize=12, spaceAfter=6)
            story.append(Paragraph(f"Clause {i}: {result.risk_level.level} Risk", clause_style))
            
            # Explanation
            story.append(Paragraph(f"<b>Explanation:</b> {result.plain_explanation}", styles['Normal']))
            
            # Recommendations
            if result.recommendations:
                rec_text = "<b>Recommendations:</b><br/>" + "<br/>".join([f"• {rec}" for rec in result.recommendations])
                story.append(Paragraph(rec_text, styles['Normal']))
            
            story.append(Spacer(1, 12))
        
        # Footer
        footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey)
        story.append(Spacer(1, 24))
        story.append(Paragraph("Generated by LegalSaathi - AI-Powered Legal Document Analysis", footer_style))
        story.append(Paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", footer_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def export_analysis_to_docx(self, analysis: DocumentAnalysis) -> BytesIO:
        """Export analysis to Word document format"""
        if not self.docx_available:
            raise ImportError("python-docx not available for Word export")
        
        doc = DocxDocument()
        
        # Title
        title = doc.add_heading('LegalSaathi Document Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(analysis.summary)
        
        # Overall Risk
        risk_para = doc.add_paragraph()
        risk_run = risk_para.add_run(f'Overall Risk Level: {analysis.overall_risk.level}')
        risk_run.bold = True
        
        # Detailed Analysis
        doc.add_heading('Detailed Clause Analysis', level=1)
        
        for i, result in enumerate(analysis.analysis_results, 1):
            # Clause header
            doc.add_heading(f'Clause {i}: {result.risk_level.level} Risk', level=2)
            
            # Explanation
            explanation_para = doc.add_paragraph()
            explanation_para.add_run('Explanation: ').bold = True
            explanation_para.add_run(result.plain_explanation)
            
            # Recommendations
            if result.recommendations:
                rec_para = doc.add_paragraph()
                rec_para.add_run('Recommendations:').bold = True
                for rec in result.recommendations:
                    doc.add_paragraph(f'• {rec}', style='List Bullet')
        
        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph('Generated by LegalSaathi - AI-Powered Legal Document Analysis')
        footer_para.add_run(f'\nReport generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

# Initialize services
service_manager = ServiceManager()
ai_orchestrator = AIServiceOrchestrator(service_manager)

# Decorators
def rate_limit_decorator(f):
    """Decorator to apply rate limiting to routes"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        client_ip = request.environ.get('HTTP_X_FORWARDED_FOR', request.remote_addr)
        rate_limiter = service_manager.get_service('rate_limiter')
        
        if rate_limiter.is_rate_limited(client_ip):
            return jsonify({
                'success': False,
                'error': 'Rate limit exceeded. Please try again in a minute.'
            }), 429
        return f(*args, **kwargs)
    return decorated_function

# Routes
@app.route('/')
def home():
    """Home page with document input interface"""
    return render_template('index.html')

@app.route('/analyze', methods=['POST'])
@rate_limit_decorator
def analyze_document():
    """Enhanced document analysis endpoint with improved error handling"""
    start_time = time.time()
    client_ip = request.environ.get('HTTP_X_FORWARDED_FOR', request.remote_addr)
    
    try:
        logger.info(f"Analysis request from IP: {client_ip}")
        
        # Get services
        file_processor = service_manager.get_service('file_processor')
        document_classifier = service_manager.get_service('document_classifier')
        expertise_detector = service_manager.get_service('expertise_detector')
        cache_manager = service_manager.get_service('cache_manager')
        metrics_tracker = service_manager.get_service('metrics_tracker')
        
        # Process input (file or text)
        document_text, file_info, warnings = process_document_input(file_processor)
        
        # Validate document input
        is_valid, error_message = DocumentValidator.validate_document_input(document_text)
        if not is_valid:
            return render_template('index.html', 
                                 error=error_message, 
                                 document_text=document_text,
                                 file_info=file_info)
        
        # Classify document type and detect expertise
        classification = document_classifier.classify_document(document_text)
        expertise_profile = detect_user_expertise(expertise_detector)
        
        # Add classification warnings
        warnings.extend(get_classification_warnings(classification))
        
        # Check cache first
        cache_key = cache_manager.get_cache_key(document_text, "analysis")
        cached_analysis = cache_manager.get_from_cache(cache_key)
        
        if cached_analysis:
            logger.info("Returning cached analysis result")
            metrics_tracker.record_analysis(True, time.time() - start_time)
            
            # Add real-time data to cached result
            cached_analysis.update({
                'processing_time': time.time() - start_time,
                'file_info': file_info,
                'warnings': warnings,
                'classification': classification,
                'expertise_profile': expertise_profile
            })
            
            return render_template('results.html', 
                                 analysis=type('Analysis', (), cached_analysis)(),
                                 file_info=file_info,
                                 warnings=warnings,
                                 classification=classification,
                                 expertise_profile=expertise_profile)
        
        # Perform AI analysis
        try:
            logger.info("Starting enhanced document analysis with Google Cloud AI")
            
            # Get document bytes for Document AI if available
            document_bytes = get_document_bytes(file_processor, file_info)
            
            # Perform analysis using AI orchestrator
            analysis = ai_orchestrator.analyze_document_with_ai(
                document_text, 
                classification.document_type, 
                expertise_profile,
                document_bytes
            )
            
            # Convert analysis to dictionary for caching
            analysis_dict = convert_analysis_to_dict(analysis)
            
            # Store in cache
            cache_manager.store_in_cache(cache_key, analysis_dict)
            
            # Record metrics
            avg_confidence = calculate_average_confidence(analysis.analysis_results)
            metrics_tracker.record_analysis(True, time.time() - start_time, avg_confidence)
            
            # Add additional data to analysis
            enhance_analysis_with_metadata(analysis, file_info, warnings, classification, expertise_profile)
            
            logger.info(f"Analysis completed successfully in {time.time() - start_time:.2f} seconds")
            
            return render_template('results.html', 
                                 analysis=analysis,
                                 file_info=file_info,
                                 warnings=warnings,
                                 classification=classification,
                                 expertise_profile=expertise_profile)
                
        except Exception as e:
            return handle_analysis_error(e, start_time, document_text, file_info, metrics_tracker)
        
    except Exception as e:
        logger.error(f"Unexpected error in analyze_document: {e}")
        return render_template('index.html', 
                             error="An unexpected error occurred. Please try again.",
                             document_text=request.form.get('document_text', ''))

# Helper functions for the main route
def process_document_input(file_processor) -> Tuple[str, Any, List[str]]:
    """Process document input from file or text"""
    document_text = ""
    file_info = None
    warnings = []
    
    # Check if file was uploaded
    if 'document_file' in request.files and request.files['document_file'].filename:
        file = request.files['document_file']
        file_result = file_processor.process_uploaded_file(file)
        
        if not file_result.success:
            raise ValueError(file_result.error_message)
        
        document_text = file_result.text_content
        file_info = file_result.file_info
        warnings.extend(file_result.warnings)
    else:
        document_text = request.form.get('document_text', '')
    
    return document_text, file_info, warnings

def detect_user_expertise(expertise_detector) -> Dict:
    """Detect user expertise level"""
    user_questions = request.form.get('user_questions', '')
    expertise_level = request.form.get('expertise_level', 'beginner')
    
    if expertise_level in ['beginner', 'intermediate', 'expert']:
        return expertise_detector._get_expertise_profile(expertise_level)
    else:
        return expertise_detector.detect_expertise(user_questions, 'medium')

def get_classification_warnings(classification) -> List[str]:
    """Get warnings based on document classification"""
    warnings = []
    classification_message = service_manager.get_service('document_classifier').get_analysis_message(classification)
    
    if classification.confidence > 0.7:
        warnings.append(classification_message)
    elif classification.confidence > 0.3:
        warnings.append(classification_message)
        warnings.append("Document type detection has moderate confidence - analysis may vary")
    else:
        warnings.append("Document type unclear - using general legal document analysis")
    
    return warnings

def get_document_bytes(file_processor, file_info) -> Optional[bytes]:
    """Get document bytes for Document AI processing"""
    if file_info and hasattr(file_processor, 'get_file_bytes'):
        try:
            return file_processor.get_file_bytes()
        except:
            pass
    return None

def convert_analysis_to_dict(analysis: DocumentAnalysis) -> Dict:
    """Convert DocumentAnalysis object to dictionary for caching"""
    return {
        'document_id': analysis.document_id,
        'analysis_results': [
            {
                'clause_id': r.clause_id,
                'risk_level': {
                    'level': r.risk_level.level,
                    'score': r.risk_level.score,
                    'reasons': r.risk_level.reasons,
                    'severity': r.risk_level.severity,
                    'confidence_percentage': r.risk_level.confidence_percentage,
                    'risk_categories': r.risk_level.risk_categories,
                    'low_confidence_warning': r.risk_level.low_confidence_warning
                },
                'plain_explanation': r.plain_explanation,
                'legal_implications': r.legal_implications,
                'recommendations': r.recommendations
            } for r in analysis.analysis_results
        ],
        'overall_risk': {
            'level': analysis.overall_risk.level,
            'score': analysis.overall_risk.score,
            'reasons': analysis.overall_risk.reasons,
            'severity': analysis.overall_risk.severity,
            'confidence_percentage': analysis.overall_risk.confidence_percentage,
            'risk_categories': analysis.overall_risk.risk_categories,
            'low_confidence_warning': analysis.overall_risk.low_confidence_warning
        },
        'summary': analysis.summary,
        'processing_time': analysis.processing_time
    }

def calculate_average_confidence(analysis_results: List[AnalysisResult]) -> float:
    """Calculate average confidence score"""
    if not analysis_results:
        return 0.0
    
    total_confidence = sum(r.risk_level.confidence_percentage for r in analysis_results)
    return total_confidence / len(analysis_results) / 100

def enhance_analysis_with_metadata(analysis, file_info, warnings, classification, expertise_profile):
    """Add metadata to analysis object"""
    analysis.file_info = file_info
    analysis.warnings = warnings
    analysis.classification = classification
    analysis.expertise_profile = expertise_profile

def handle_analysis_error(error: Exception, start_time: float, document_text: str, file_info: Any, metrics_tracker) -> str:
    """Handle analysis errors with appropriate user messages"""
    error_type = type(error).__name__
    processing_time = time.time() - start_time
    
    logger.error(f"Analysis error: {error}")
    
    if isinstance(error, ConnectionError):
        metrics_tracker.record_analysis(False, processing_time, error_type="ConnectionError")
        error_msg = "Unable to connect to AI analysis service. Please try again in a few moments."
    elif isinstance(error, TimeoutError):
        metrics_tracker.record_analysis(False, processing_time, error_type="TimeoutError")
        error_msg = "Analysis service timed out. Please try again with a shorter document or try again later."
    elif "memory" in str(error).lower() or "resource" in str(error).lower():
        metrics_tracker.record_analysis(False, processing_time, error_type="MemoryError")
        error_msg = "Analysis service is currently overloaded. Please try again in a few minutes."
    else:
        metrics_tracker.record_analysis(False, processing_time, error_type=error_type)
        error_msg = "An error occurred during analysis. Please try again or contact support if the problem persists."
    
    return render_template('index.html', 
                         error=error_msg, 
                         document_text=document_text,
                         file_info=file_info)

# Additional API routes
@app.route('/api/clarify', methods=['POST'])
@rate_limit_decorator
def clarify_question():
    """AI-powered clarification endpoint"""
    try:
        data = request.get_json()
        question = data.get('question', '')
        context = data.get('context', {})
        
        if not question.strip():
            return jsonify({
                'success': False,
                'error': 'Question is required'
            }), 400
        
        ai_service = service_manager.get_service('ai_clarification')
        result = ai_service.ask_clarification(question, context)
        
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Clarification error: {e}")
        return jsonify({
            'success': False,
            'error': 'Unable to process clarification request'
        }), 500

@app.route('/api/export/pdf/<document_id>')
def export_pdf(document_id):
    """Export analysis as PDF"""
    try:
        # Get analysis from cache or database
        cache_manager = service_manager.get_service('cache_manager')
        analysis_data = cache_manager.get_from_cache(f"analysis:{document_id}")
        
        if not analysis_data:
            return jsonify({'error': 'Analysis not found'}), 404
        
        # Convert to DocumentAnalysis object
        analysis = convert_dict_to_analysis(analysis_data)
        
        # Generate PDF
        report_exporter = service_manager.get_service('report_exporter')
        pdf_buffer = report_exporter.export_analysis_to_pdf(analysis)
        
        return send_file(
            pdf_buffer,
            as_attachment=True,
            download_name=f'legal_analysis_{document_id}.pdf',
            mimetype='application/pdf'
        )
        
    except Exception as e:
        logger.error(f"PDF export error: {e}")
        return jsonify({'error': 'Export failed'}), 500

@app.route('/api/export/docx/<document_id>')
def export_docx(document_id):
    """Export analysis as Word document"""
    try:
        # Get analysis from cache or database
        cache_manager = service_manager.get_service('cache_manager')
        analysis_data = cache_manager.get_from_cache(f"analysis:{document_id}")
        
        if not analysis_data:
            return jsonify({'error': 'Analysis not found'}), 404
        
        # Convert to DocumentAnalysis object
        analysis = convert_dict_to_analysis(analysis_data)
        
        # Generate Word document
        report_exporter = service_manager.get_service('report_exporter')
        docx_buffer = report_exporter.export_analysis_to_docx(analysis)
        
        return send_file(
            docx_buffer,
            as_attachment=True,
            download_name=f'legal_analysis_{document_id}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"Word export error: {e}")
        return jsonify({'error': 'Export failed'}), 500

@app.route('/health')
def health_check():
    """System health check endpoint"""
    try:
        health_status = {
            'status': 'healthy',
            'timestamp': datetime.utcnow().isoformat(),
            'services': {
                'google_cloud_ai': 'operational' if service_manager.get_service('google_document_ai').enabled else 'degraded',
                'natural_language': 'operational' if service_manager.get_service('google_natural_language').enabled else 'degraded',
                'speech_service': 'operational' if service_manager.get_service('google_speech').enabled else 'degraded',
                'translation': 'operational',
                'ai_clarification': 'operational'
            },
            'performance': service_manager.get_service('metrics_tracker').get_metrics()
        }
        
        return jsonify(health_status)
        
    except Exception as e:
        logger.error(f"Health check error: {e}")
        return jsonify({
            'status': 'unhealthy',
            'error': str(e)
        }), 500

def convert_dict_to_analysis(analysis_data: Dict) -> DocumentAnalysis:
    """Convert dictionary back to DocumentAnalysis object"""
    # Convert analysis results
    analysis_results = []
    for result_data in analysis_data['analysis_results']:
        risk_level = RiskLevel(**result_data['risk_level'])
        analysis_result = AnalysisResult(
            clause_id=result_data['clause_id'],
            risk_level=risk_level,
            plain_explanation=result_data['plain_explanation'],
            legal_implications=result_data['legal_implications'],
            recommendations=result_data['recommendations']
        )
        analysis_results.append(analysis_result)
    
    # Convert overall risk
    overall_risk = RiskLevel(**analysis_data['overall_risk'])
    
    return DocumentAnalysis(
        document_id=analysis_data['document_id'],
        analysis_results=analysis_results,
        overall_risk=overall_risk,
        summary=analysis_data['summary'],
        processing_time=analysis_data['processing_time']
    )

# Application startup
if __name__ == '__main__':
    logger.info("LegalSaathi Document Advisor starting up...")
    logger.info(f"Available export formats: {['pdf' if REPORTLAB_AVAILABLE else None, 'docx' if PYTHON_DOCX_AVAILABLE else None]}")
    logger.info(f"Cache expiry: {config.CACHE_EXPIRY} seconds")
    logger.info(f"Rate limit: {config.RATE_LIMIT_REQUESTS} requests per {config.RATE_LIMIT_WINDOW} seconds")
    logger.info("Application initialized successfully")
    
    app.run(debug=True, host='0.0.0.0', port=5000)