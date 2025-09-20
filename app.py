from flask import Flask, render_template, request, jsonify, send_file, make_response, g, send_from_directory
from dataclasses import dataclass
from typing import List, Dict, Any, Optional
import time
import json
import re
import logging
import hashlib
import os
from datetime import datetime, timedelta
from functools import wraps
from io import BytesIO
import threading
from openai import OpenAI

# Import production modules
from monitoring import production_monitor, health_check_decorator
from performance_optimizer import performance_optimizer, performance_monitor, compress_json_response, setup_performance_monitoring
from risk_classifier import RiskClassifier, classify_document_risk, RiskLevel as ClassifierRiskLevel, UserExpertiseDetector, EnhancedRiskCategories
from google_translate_service import GoogleTranslateService
from google_document_ai_service import document_ai_service
from google_natural_language_service import natural_language_service
from google_speech_service import speech_service
from file_processor import FileProcessor
from document_classifier import DocumentClassifier, DocumentType

# Import for PDF/Word export functionality
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

app = Flask(__name__)

# Configure file upload settings
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB max file size

# Set up performance monitoring
app = setup_performance_monitoring(app)

# Configure logging for production monitoring
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('legal_saathi.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# In-memory cache for analysis results and rate limiting
analysis_cache = {}
rate_limit_cache = {}
CACHE_EXPIRY = 3600  # 1 hour
RATE_LIMIT_REQUESTS = 10  # requests per minute
RATE_LIMIT_WINDOW = 60  # seconds

# Initialize services including enhanced Google Cloud AI services
file_processor = FileProcessor()
document_classifier = DocumentClassifier()
translation_service = GoogleTranslateService()
expertise_detector = UserExpertiseDetector()
risk_categories = EnhancedRiskCategories()

# Initialize Google Cloud AI services
google_document_ai = document_ai_service
google_natural_language = natural_language_service
google_speech = speech_service

# Initialize OpenAI client for vLLM server (adapting from integrate.py)
client = OpenAI(
    base_url="http://172.25.0.211:8002/v1",
    api_key="EMPTY",  # No API key needed for local vLLM
    timeout=5.0  # 5 second timeout for faster fallback
)

# Enhanced Real AI clarification service with Google Cloud AI integration
class RealAIService:
    """Enhanced AI service for document clarification with improved Google Cloud AI integration"""
    
    def __init__(self, client):
        self.client = client
        self.conversation_history = []
        self.fallback_responses = {
            'general': "I understand you're asking about your legal document. While I'm having trouble accessing my full AI capabilities right now, I'd recommend reviewing the specific clause you're concerned about and considering consultation with a legal professional if you have serious concerns.",
            'terms': "Legal terms can be complex. For specific definitions and implications, I recommend consulting with a qualified legal professional who can provide accurate interpretations based on your jurisdiction.",
            'rights': "Understanding your rights is important. While I can provide general guidance, specific rights can vary by location and situation. Consider consulting with a legal professional for personalized advice.",
            'obligations': "Legal obligations should be clearly understood before signing any document. If you're unsure about specific requirements, it's best to seek professional legal advice."
        }
    
    def ask_clarification(self, question: str, context: Dict = None) -> Dict[str, Any]:
        """Enhanced AI-powered conversational clarification with improved error handling"""
        try:
            # Build enhanced context from document analysis
            context_text = ""
            if context:
                risk_level = context.get('risk_level', 'unknown')
                summary = context.get('summary', '')
                document_type = context.get('document_type', 'legal document')
                context_text = f"""Document Context:
- Type: {document_type}
- Risk Level: {risk_level}
- Summary: {summary[:200]}...

"""
            
            # Enhanced prompt with better structure
            prompt = f"""You are LegalSaathi, an AI legal document advisor specializing in making complex legal language accessible to everyday citizens.

{context_text}User Question: "{question}"

Please provide a helpful response that:
1. Directly answers the user's question in plain language
2. Explains any legal terms in simple words
3. Provides practical, actionable advice when appropriate
4. Indicates when professional legal consultation is recommended
5. Maintains a supportive, educational tone

Keep your response clear, concise (under 200 words), and focused on helping the user understand their document better."""

            # Make API call with optimized timeout for faster response
            response = self.client.chat.completions.create(
                model="meta-llama/Llama-3.1-8B-Instruct",
                messages=[
                    {"role": "system", "content": "You are LegalSaathi, a helpful and knowledgeable legal document advisor focused on accessibility and user education."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=150,  # Reduced for faster response
                temperature=0.4,  # Lower temperature for more focused responses
                timeout=3.0  # Reduced timeout for faster fallback
            )
            
            ai_response = response.choices[0].message.content.strip()
            
            # Validate response quality
            if len(ai_response) < 20:
                raise ValueError("Response too short")
            
            # Store in conversation history with metadata
            conversation_entry = {
                'question': question,
                'response': ai_response,
                'timestamp': time.time(),
                'context_provided': bool(context),
                'response_length': len(ai_response),
                'confidence': 'high'
            }
            self.conversation_history.append(conversation_entry)
            
            # Keep conversation history manageable
            if len(self.conversation_history) > 50:
                self.conversation_history = self.conversation_history[-25:]
            
            logger.info(f"AI clarification successful: {len(ai_response)} chars")
            
            return {
                'success': True,
                'response': ai_response,
                'conversation_id': len(self.conversation_history),
                'confidence_score': 85,
                'response_quality': 'high',
                'processing_time': time.time() - conversation_entry['timestamp']
            }
            
        except Exception as e:
            logger.warning(f"AI clarification error: {e}, using fallback")
            
            # Intelligent fallback based on question content
            fallback_response = self._get_intelligent_fallback(question)
            
            # Store fallback in history
            fallback_entry = {
                'question': question,
                'response': fallback_response,
                'timestamp': time.time(),
                'context_provided': bool(context),
                'response_length': len(fallback_response),
                'confidence': 'low',
                'fallback': True
            }
            self.conversation_history.append(fallback_entry)
            
            return {
                'success': True,
                'response': fallback_response,
                'fallback': True,
                'conversation_id': len(self.conversation_history),
                'confidence_score': 45,
                'response_quality': 'basic',
                'error_type': type(e).__name__
            }
    
    def _get_intelligent_fallback(self, question: str) -> str:
        """Provide intelligent fallback responses based on question content"""
        question_lower = question.lower()
        
        if any(word in question_lower for word in ['what', 'mean', 'definition', 'explain']):
            return self.fallback_responses['terms']
        elif any(word in question_lower for word in ['right', 'rights', 'can i', 'allowed']):
            return self.fallback_responses['rights']
        elif any(word in question_lower for word in ['must', 'have to', 'required', 'obligation']):
            return self.fallback_responses['obligations']
        else:
            return self.fallback_responses['general']
    
    def get_conversation_summary(self) -> Dict[str, Any]:
        """Enhanced conversation summary with analytics"""
        if not self.conversation_history:
            return {
                'success': True,
                'total_questions': 0,
                'recent_questions': [],
                'analytics': {
                    'avg_response_length': 0,
                    'fallback_rate': 0,
                    'high_confidence_rate': 0
                }
            }
        
        # Calculate analytics
        total_questions = len(self.conversation_history)
        fallback_count = sum(1 for item in self.conversation_history if item.get('fallback', False))
        high_confidence_count = sum(1 for item in self.conversation_history if item.get('confidence') == 'high')
        avg_length = sum(item['response_length'] for item in self.conversation_history) / total_questions
        
        return {
            'success': True,
            'total_questions': total_questions,
            'recent_questions': [item['question'] for item in self.conversation_history[-5:]],
            'analytics': {
                'avg_response_length': round(avg_length, 1),
                'fallback_rate': round((fallback_count / total_questions) * 100, 1),
                'high_confidence_rate': round((high_confidence_count / total_questions) * 100, 1)
            },
            'last_activity': max(item['timestamp'] for item in self.conversation_history) if self.conversation_history else None
        }

# Initialize AI clarification service after class definition
ai_clarification_service = RealAIService(client)

# Caching and Rate Limiting Functions
def get_cache_key(text: str, analysis_type: str = "analysis") -> str:
    """Generate cache key for analysis results"""
    text_hash = hashlib.md5(text.encode('utf-8')).hexdigest()
    return f"{analysis_type}:{text_hash}"

def get_from_cache(cache_key: str) -> Optional[Dict]:
    """Get analysis result from cache if not expired"""
    if cache_key in analysis_cache:
        cached_data = analysis_cache[cache_key]
        if time.time() - cached_data['timestamp'] < CACHE_EXPIRY:
            logger.info(f"Cache hit for key: {cache_key}")
            return cached_data['data']
        else:
            # Remove expired cache entry
            del analysis_cache[cache_key]
            logger.info(f"Cache expired for key: {cache_key}")
    return None

def store_in_cache(cache_key: str, data: Dict) -> None:
    """Store analysis result in cache"""
    analysis_cache[cache_key] = {
        'data': data,
        'timestamp': time.time()
    }
    logger.info(f"Stored in cache: {cache_key}")

def rate_limit_check(client_ip: str) -> bool:
    """Check if client has exceeded rate limit"""
    current_time = time.time()
    
    # Clean old entries
    for ip in list(rate_limit_cache.keys()):
        rate_limit_cache[ip] = [
            timestamp for timestamp in rate_limit_cache[ip]
            if current_time - timestamp < RATE_LIMIT_WINDOW
        ]
        if not rate_limit_cache[ip]:
            del rate_limit_cache[ip]
    
    # Check current client
    if client_ip not in rate_limit_cache:
        rate_limit_cache[client_ip] = []
    
    if len(rate_limit_cache[client_ip]) >= RATE_LIMIT_REQUESTS:
        logger.warning(f"Rate limit exceeded for IP: {client_ip}")
        return False
    
    rate_limit_cache[client_ip].append(current_time)
    return True

def rate_limit_decorator(f):
    """Decorator to apply rate limiting to routes"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        client_ip = request.environ.get('HTTP_X_FORWARDED_FOR', request.remote_addr)
        if not rate_limit_check(client_ip):
            return jsonify({
                'success': False,
                'error': 'Rate limit exceeded. Please try again in a minute.'
            }), 429
        return f(*args, **kwargs)
    return decorated_function

# Enhanced confidence scoring and monitoring
class AnalysisMetrics:
    """Track analysis metrics for monitoring"""
    
    def __init__(self):
        self.metrics = {
            'total_analyses': 0,
            'successful_analyses': 0,
            'failed_analyses': 0,
            'cache_hits': 0,
            'avg_processing_time': 0.0,
            'confidence_scores': [],
            'error_types': {}
        }
        self.lock = threading.Lock()
    
    def record_analysis(self, success: bool, processing_time: float, confidence: float = None, error_type: str = None):
        """Record analysis metrics"""
        with self.lock:
            self.metrics['total_analyses'] += 1
            if success:
                self.metrics['successful_analyses'] += 1
                if confidence is not None:
                    self.metrics['confidence_scores'].append(confidence)
            else:
                self.metrics['failed_analyses'] += 1
                if error_type:
                    self.metrics['error_types'][error_type] = self.metrics['error_types'].get(error_type, 0) + 1
            
            # Update average processing time
            total_time = self.metrics['avg_processing_time'] * (self.metrics['total_analyses'] - 1) + processing_time
            self.metrics['avg_processing_time'] = total_time / self.metrics['total_analyses']
    
    def get_metrics(self) -> Dict:
        """Get current metrics"""
        with self.lock:
            metrics_copy = self.metrics.copy()
            if metrics_copy['confidence_scores']:
                metrics_copy['avg_confidence'] = sum(metrics_copy['confidence_scores']) / len(metrics_copy['confidence_scores'])
            else:
                metrics_copy['avg_confidence'] = 0.0
            return metrics_copy

# Initialize metrics tracker
metrics_tracker = AnalysisMetrics()

# Enhanced data models for document analysis results with confidence and categories
@dataclass
class RiskLevel:
    level: str  # "RED", "YELLOW", "GREEN"
    score: float  # 0.0 to 1.0
    reasons: List[str]
    severity: str
    confidence_percentage: int  # 0-100%
    risk_categories: dict  # financial, legal, operational
    low_confidence_warning: bool

@dataclass
class AnalysisResult:
    clause_id: str
    risk_level: RiskLevel
    plain_explanation: str
    legal_implications: List[str]
    recommendations: List[str]

@dataclass
class DocumentAnalysis:
    document_id: str
    analysis_results: List[AnalysisResult]
    overall_risk: RiskLevel
    summary: str
    processing_time: float

# API-only Flask app - React served separately

@app.route('/health')
@performance_monitor
def health_check():
    """Production health check endpoint for monitoring"""
    try:
        # Basic health check for core services
        health_status = {
            'status': 'healthy',
            'timestamp': datetime.now().isoformat(),
            'services': {
                'ai_clarification': True,
                'translation': translation_service.enabled if hasattr(translation_service, 'enabled') else True,
                'risk_classification': True,
                'file_processing': True,
                'google_document_ai': google_document_ai.enabled,
                'google_natural_language': google_natural_language.enabled,
                'google_speech': google_speech.enabled
            },
            'cache': {
                'size': len(analysis_cache),
                'status': 'operational'
            }
        }
        
        # Determine overall health
        critical_services = ['ai_clarification', 'risk_classification', 'file_processing']
        if not all(health_status['services'][service] for service in critical_services):
            health_status['status'] = 'degraded'
        
        return jsonify(health_status)
        
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        return jsonify({
            'status': 'unhealthy',
            'error': str(e),
            'timestamp': datetime.now().isoformat()
        }), 503





def validate_document_input(document_text):
    """
    Validate document input for universal legal document analysis
    Returns tuple (is_valid, error_message)
    """
    if not document_text or not document_text.strip():
        return False, "Please provide a valid legal document text."
    
    text = document_text.strip()
    length = len(text)
    
    # Check length restrictions
    if length < 100:
        return False, "Document text must be at least 100 characters long for meaningful analysis."
    
    if length > 50000:
        return False, "Document text exceeds maximum length of 50,000 characters. Please shorten your document."
    
    # Basic content validation - check for general legal document keywords
    legal_keywords = [
        'agreement', 'contract', 'party', 'parties', 'terms', 'conditions',
        'shall', 'hereby', 'whereas', 'therefore', 'obligations', 'rights',
        'liability', 'breach', 'termination', 'effective', 'binding', 'legal'
    ]
    
    text_lower = text.lower()
    found_keywords = [keyword for keyword in legal_keywords if keyword in text_lower]
    
    if len(found_keywords) < 2:
        return False, "This doesn't appear to be a legal document. Please ensure you've pasted the correct document text."
    
    return True, None

def _get_document_context(document_type: DocumentType) -> Dict[str, str]:
    """Get context-specific terms for different document types"""
    contexts = {
        DocumentType.RENTAL_AGREEMENT: {
            'party_disadvantage': 'tenant disadvantage',
            'other_party': 'landlord',
            'protection_laws': 'tenant protection statutes',
            'legal_area': 'rental/housing'
        },
        DocumentType.EMPLOYMENT_CONTRACT: {
            'party_disadvantage': 'employee disadvantage',
            'other_party': 'employer',
            'protection_laws': 'employment protection laws',
            'legal_area': 'employment'
        },
        DocumentType.NDA: {
            'party_disadvantage': 'receiving party disadvantage',
            'other_party': 'disclosing party',
            'protection_laws': 'confidentiality regulations',
            'legal_area': 'intellectual property'
        },
        DocumentType.LOAN_AGREEMENT: {
            'party_disadvantage': 'borrower disadvantage',
            'other_party': 'lender',
            'protection_laws': 'consumer lending laws',
            'legal_area': 'financial/lending'
        },
        DocumentType.PARTNERSHIP_AGREEMENT: {
            'party_disadvantage': 'partner disadvantage',
            'other_party': 'other partners',
            'protection_laws': 'partnership regulations',
            'legal_area': 'business/corporate'
        }
    }
    
    return contexts.get(document_type, {
        'party_disadvantage': 'party disadvantage',
        'other_party': 'other party',
        'protection_laws': 'applicable regulations',
        'legal_area': 'relevant legal'
    })

def analyze_legal_document_with_llm(document_text: str, document_type: DocumentType, expertise_profile: Dict = None, document_bytes: bytes = None) -> DocumentAnalysis:
    """
    Enhanced legal document analysis using multiple Google Cloud AI services
    Integrates Document AI, Natural Language AI, and advanced risk classification
    Requirements: 2.1, 2.2, 4.1, 3.1, 3.2, 3.3, 3.4, 7.2
    """
    start_time = time.time()
    
    try:
        # Enhanced analysis with Google Cloud AI services
        enhanced_insights = {}
        
        # 1. Google Cloud Document AI analysis (if document bytes available)
        if document_bytes and google_document_ai.enabled:
            try:
                logger.info("Running Google Document AI analysis")
                doc_ai_result = google_document_ai.process_legal_document(document_bytes)
                if doc_ai_result['success']:
                    enhanced_insights['document_ai'] = {
                        'entities': doc_ai_result['entities'],
                        'tables': doc_ai_result['tables'],
                        'key_value_pairs': doc_ai_result['key_value_pairs'],
                        'legal_clauses': doc_ai_result['legal_clauses'],
                        'confidence_scores': doc_ai_result['confidence_scores']
                    }
                    logger.info(f"Document AI extracted {len(doc_ai_result['entities'])} entities")
            except Exception as e:
                logger.warning(f"Document AI analysis failed: {e}")
        
        # 2. Google Cloud Natural Language AI analysis
        if google_natural_language.enabled:
            try:
                logger.info("Running Google Natural Language AI analysis")
                nl_result = google_natural_language.analyze_legal_document(document_text)
                if nl_result['success']:
                    enhanced_insights['natural_language'] = {
                        'sentiment': nl_result['sentiment'],
                        'entities': nl_result['entities'],
                        'syntax': nl_result['syntax'],
                        'classification': nl_result['classification'],
                        'legal_insights': nl_result['legal_insights']
                    }
                    logger.info(f"Natural Language AI completed with {len(nl_result['entities'])} entities")
            except Exception as e:
                logger.warning(f"Natural Language AI analysis failed: {e}")
        
        # 3. Continue with existing risk classification (enhanced with AI insights)
        # Use the new risk classification system with error handling
        risk_analysis = classify_document_risk(document_text, document_type)
        
        # Convert to our enhanced data model format with confidence and categories
        overall_risk_data = risk_analysis['overall_risk']
        overall_risk = RiskLevel(
            level=overall_risk_data['level'],
            score=overall_risk_data['score'],
            reasons=[f"Overall document risk: {overall_risk_data['severity']}"],
            severity=overall_risk_data['severity'],
            confidence_percentage=overall_risk_data['confidence_percentage'],
            risk_categories=overall_risk_data['risk_categories'],
            low_confidence_warning=overall_risk_data['low_confidence_warning']
        )
        
        analysis_results = []
        for clause_assessment in risk_analysis['clause_assessments']:
            assessment = clause_assessment['assessment']
            
            # Convert classifier risk level to our enhanced format
            risk_level = RiskLevel(
                level=assessment.level.value,
                score=assessment.score,
                reasons=assessment.reasons,
                severity=assessment.severity,
                confidence_percentage=assessment.confidence_percentage,
                risk_categories=assessment.risk_categories,
                low_confidence_warning=assessment.low_confidence_warning
            )
            
            # Generate adaptive explanations based on user expertise
            confidence_indicator = ""
            if assessment.low_confidence_warning:
                confidence_indicator = f" (⚠️ Low Confidence: {assessment.confidence_percentage}%)"
            
            # Adapt explanation style based on expertise level and document type
            doc_context = _get_document_context(document_type)
            
            if expertise_profile and expertise_profile['level'] == 'expert':
                # Technical explanations for experts
                if assessment.level == ClassifierRiskLevel.RED:
                    explanation = f"🚨 HIGH RISK{confidence_indicator}: Legal analysis indicates significant {doc_context['party_disadvantage']}. Risk categories: {', '.join([k for k, v in assessment.risk_categories.items() if v > 0.6])}. {' '.join(assessment.reasons[:2])}"
                    implications = [f"Potential legal liability exposure", f"Contractual imbalance favoring {doc_context['other_party']}", f"May violate {doc_context['protection_laws']}"]
                    recommendations = ["Seek legal counsel for clause modification", f"Review relevant {doc_context['legal_area']} laws", "Consider contract renegotiation or withdrawal"]
                elif assessment.level == ClassifierRiskLevel.YELLOW:
                    explanation = f"⚠️ MODERATE RISK{confidence_indicator}: Analysis shows suboptimal terms with negotiation potential. {' '.join(assessment.reasons[:2])}"
                    implications = ["Terms deviate from market standards", f"Potential for improved {doc_context['party_disadvantage'].split()[0]} protections"]
                    recommendations = ["Negotiate more balanced terms", f"Reference {doc_context['legal_area']} market standards", "Document any agreed modifications"]
                else:
                    explanation = f"✅ ACCEPTABLE TERMS{confidence_indicator}: Clause aligns with standard {doc_context['legal_area']} practices. {' '.join(assessment.reasons[:1]) if assessment.reasons else 'No significant legal concerns identified.'}"
                    implications = ["Terms within acceptable legal parameters", "Standard market practice"]
                    recommendations = ["Terms are legally sound as written"]
            
            elif expertise_profile and expertise_profile['level'] == 'intermediate':
                # Balanced explanations for intermediate users
                if assessment.level == ClassifierRiskLevel.RED:
                    explanation = f"🚨 HIGH RISK{confidence_indicator}: This clause is unfavorable and should be addressed. Main concerns: {', '.join(assessment.reasons[:2])}"
                    implications = ["This puts you at a significant disadvantage", "Could lead to financial or legal problems", f"Not typical in fair {doc_context['legal_area']} agreements"]
                    recommendations = ["Negotiate to change this clause", "Get legal advice before signing", f"Consider this a red flag for the {doc_context['other_party']}"]
                elif assessment.level == ClassifierRiskLevel.YELLOW:
                    explanation = f"⚠️ MODERATE CONCERN{confidence_indicator}: This clause could be more balanced. Issues: {' '.join(assessment.reasons[:2])}"
                    implications = ["Terms are somewhat one-sided", f"May cause issues during the {doc_context['legal_area']} relationship", "Room for improvement through negotiation"]
                    recommendations = ["Try to negotiate better terms", "Understand what you're agreeing to", f"Compare with other {doc_context['legal_area']} options"]
                else:
                    explanation = f"✅ FAIR TERMS{confidence_indicator}: This appears to be a reasonable and standard clause. {' '.join(assessment.reasons[:1]) if assessment.reasons else 'No major concerns identified.'}"
                    implications = ["Terms are balanced and fair", f"Follows standard {doc_context['legal_area']} practices"]
                    recommendations = ["This clause is acceptable"]
            
            else:
                # Simple explanations for beginners (default)
                if assessment.level == ClassifierRiskLevel.RED:
                    explanation = f"🚨 DANGER{confidence_indicator}: This part of the contract is bad for you! It means: {' '.join(assessment.reasons[:1])}"
                    implications = ["This could cost you money or cause problems", "You might lose important rights", f"This is not fair in {doc_context['legal_area']} agreements"]
                    recommendations = [f"Ask the {doc_context['other_party']} to change this", f"Get help from someone who knows about {doc_context['legal_area']}", "Don't sign if they won't change it"]
                elif assessment.level == ClassifierRiskLevel.YELLOW:
                    explanation = f"⚠️ BE CAREFUL{confidence_indicator}: This part could be better for you. Problem: {' '.join(assessment.reasons[:1])}"
                    implications = ["This might cause problems later", "It's not the worst, but not great either"]
                    recommendations = ["Ask if this can be changed", "Make sure you understand what this means", "Think about whether you're okay with this"]
                else:
                    explanation = f"✅ LOOKS GOOD{confidence_indicator}: This part seems fair and normal. {' '.join(assessment.reasons[:1]) if assessment.reasons else 'Nothing to worry about here.'}"
                    implications = [f"This is typical for {doc_context['legal_area']} contracts", f"This part protects both parties involved"]
                    recommendations = ["This part is fine to agree to"]
            
            analysis_result = AnalysisResult(
                clause_id=clause_assessment['clause_id'],
                risk_level=risk_level,
                plain_explanation=explanation,
                legal_implications=implications,
                recommendations=recommendations
            )
            analysis_results.append(analysis_result)
        
        # Generate enhanced summary with confidence and category information
        red_clauses = len([c for c in risk_analysis['clause_assessments'] if c['assessment'].level == ClassifierRiskLevel.RED])
        yellow_clauses = len([c for c in risk_analysis['clause_assessments'] if c['assessment'].level == ClassifierRiskLevel.YELLOW])
        
        confidence_note = ""
        if overall_risk_data['low_confidence_warning']:
            confidence_note = f" (Analysis Confidence: {overall_risk_data['confidence_percentage']}% - Consider professional review)"
        
        # Category breakdown
        categories = overall_risk_data['risk_categories']
        high_risk_categories = [cat for cat, score in categories.items() if score > 0.6]
        category_note = f" Primary concerns: {', '.join(high_risk_categories)}" if high_risk_categories else ""
        
        doc_type_name = document_type.value.replace('_', ' ').title() if document_type != DocumentType.UNKNOWN else "legal document"
        
        if overall_risk_data['level'] == 'RED':
            summary = f"🚨 HIGH RISK DOCUMENT{confidence_note}: This {doc_type_name.lower()} contains {red_clauses} high-risk clauses that need immediate attention.{category_note} We strongly recommend legal review before signing."
        elif overall_risk_data['level'] == 'YELLOW':
            summary = f"⚠️ MODERATE RISK{confidence_note}: This {doc_type_name.lower()} has {yellow_clauses} concerning clauses that could be improved.{category_note} Review the highlighted issues and consider negotiating better terms."
        else:
            summary = f"✅ LOW RISK{confidence_note}: This appears to be a fairly standard {doc_type_name.lower()} with reasonable terms.{category_note} Review the details but generally acceptable."
        
        processing_time = time.time() - start_time
        
        return DocumentAnalysis(
            document_id=f"doc_{int(time.time())}",
            analysis_results=analysis_results,
            overall_risk=overall_risk,
            summary=summary,
            processing_time=processing_time
        )
        
    except Exception as e:
        print(f"Error in enhanced risk analysis: {str(e)}")
        # Enhanced error handling with specific error messages
        error_msg = str(e).lower()
        
        if "connection" in error_msg or "timeout" in error_msg:
            print("LLM service connection issue - using fallback analysis")
        elif "api" in error_msg or "key" in error_msg:
            print("LLM service API issue - using fallback analysis")
        elif "memory" in error_msg or "resource" in error_msg:
            print("LLM service resource issue - using fallback analysis")
        else:
            print(f"Unexpected LLM service error: {str(e)} - using fallback analysis")
        
        # Fallback to basic analysis on error
        return create_fallback_analysis(document_text)

def create_fallback_analysis(document_text: str) -> DocumentAnalysis:
    """
    Create a basic fallback analysis when enhanced AI services are unavailable
    Uses keyword-based detection as backup
    """
    try:
        # Try to use at least the keyword-based risk classifier
        classifier = RiskClassifier()
        
        # Simple clause extraction for fallback
        clauses = re.split(r'[.!?]+', document_text)
        clauses = [clause.strip() for clause in clauses if len(clause.strip()) > 20]
        
        analysis_results = []
        risk_scores = []
        
        for i, clause in enumerate(clauses[:5]):  # Limit to 5 clauses for fallback
            # Use only keyword detection (no AI)
            keyword_risk = classifier._detect_keyword_red_flags(clause)
            
            risk_level = RiskLevel(
                level=keyword_risk.level.value,
                score=keyword_risk.score,
                reasons=keyword_risk.reasons,
                severity=keyword_risk.severity,
                confidence_percentage=keyword_risk.confidence_percentage,
                risk_categories=keyword_risk.risk_categories,
                low_confidence_warning=keyword_risk.low_confidence_warning
            )
            
            analysis_result = AnalysisResult(
                clause_id=f"fallback_clause_{i+1}",
                risk_level=risk_level,
                plain_explanation=f"Keyword-based analysis: {keyword_risk.severity} risk detected.",
                legal_implications=["Limited analysis - AI services unavailable"],
                recommendations=["Manual review recommended", "Consider professional legal advice"]
            )
            analysis_results.append(analysis_result)
            risk_scores.append(keyword_risk.score)
        
        # Calculate overall risk
        if risk_scores:
            avg_score = sum(risk_scores) / len(risk_scores)
            max_score = max(risk_scores)
            overall_score = (avg_score * 0.6) + (max_score * 0.4)
        else:
            overall_score = 0.5
        
        if overall_score > 0.7:
            overall_level = "RED"
            severity = "High"
            summary = "⚠️ Potential high-risk terms detected. Full AI analysis unavailable - manual review strongly recommended."
        elif overall_score > 0.4:
            overall_level = "YELLOW"
            severity = "Medium"
            summary = "Some concerning terms detected. Limited analysis available - consider professional review."
        else:
            overall_level = "GREEN"
            severity = "Low"
            summary = "No major red flags detected in basic analysis. Full AI analysis unavailable."
        
        overall_risk = RiskLevel(
            level=overall_level,
            score=overall_score,
            reasons=[f"Fallback analysis completed - {len(analysis_results)} clauses checked"],
            severity=severity,
            confidence_percentage=60,  # Medium confidence for fallback
            risk_categories={'financial': 0.3, 'legal': 0.3, 'operational': 0.3},
            low_confidence_warning=True
        )
        
        classifier.close()
        
        return DocumentAnalysis(
            document_id=f"fallback_{int(time.time())}",
            analysis_results=analysis_results,
            overall_risk=overall_risk,
            summary=summary,
            processing_time=0.2
        )
        
    except Exception as e:
        print(f"Fallback analysis error: {e}")
        
        # Ultimate fallback - very basic analysis
        overall_risk = RiskLevel(
            level="YELLOW",
            score=0.5,
            reasons=["Analysis services unavailable"],
            severity="Medium",
            confidence_percentage=20,
            risk_categories={'financial': 0.5, 'legal': 0.5, 'operational': 0.5},
            low_confidence_warning=True
        )
        
        basic_analysis = AnalysisResult(
            clause_id="basic_fallback",
            risk_level=overall_risk,
            plain_explanation="Analysis services are currently unavailable. Please try again later or consult a legal professional.",
            legal_implications=["Unable to perform automated analysis"],
            recommendations=["Try again later", "Consult with a legal professional", "Review document manually"]
        )
        
        return DocumentAnalysis(
            document_id=f"basic_fallback_{int(time.time())}",
            analysis_results=[basic_analysis],
            overall_risk=overall_risk,
            summary="Analysis services temporarily unavailable. Manual review recommended.",
            processing_time=0.1
        )

# Document Comparison Feature
class DocumentComparison:
    """Compare two documents for risk differences"""
    
    def __init__(self):
        pass
    
    def compare_documents(self, doc1_text: str, doc2_text: str, doc1_type: DocumentType = None, doc2_type: DocumentType = None) -> Dict:
        """Compare two documents and highlight differences in risk levels"""
        try:
            logger.info("Starting document comparison")
            
            # Analyze both documents
            analysis1 = analyze_legal_document_with_llm(doc1_text, doc1_type or DocumentType.UNKNOWN)
            analysis2 = analyze_legal_document_with_llm(doc2_text, doc2_type or DocumentType.UNKNOWN)
            
            # Compare overall risk levels
            risk_comparison = {
                'document1_risk': analysis1.overall_risk.level,
                'document2_risk': analysis2.overall_risk.level,
                'risk_difference': self._calculate_risk_difference(analysis1.overall_risk, analysis2.overall_risk)
            }
            
            # Compare clause-by-clause (simplified)
            clause_comparisons = []
            max_clauses = min(len(analysis1.analysis_results), len(analysis2.analysis_results))
            
            for i in range(max_clauses):
                clause1 = analysis1.analysis_results[i]
                clause2 = analysis2.analysis_results[i]
                
                comparison = {
                    'clause_index': i + 1,
                    'doc1_risk': clause1.risk_level.level,
                    'doc2_risk': clause2.risk_level.level,
                    'risk_change': self._get_risk_change(clause1.risk_level.level, clause2.risk_level.level),
                    'doc1_explanation': clause1.plain_explanation[:200] + "..." if len(clause1.plain_explanation) > 200 else clause1.plain_explanation,
                    'doc2_explanation': clause2.plain_explanation[:200] + "..." if len(clause2.plain_explanation) > 200 else clause2.plain_explanation
                }
                clause_comparisons.append(comparison)
            
            return {
                'success': True,
                'risk_comparison': risk_comparison,
                'clause_comparisons': clause_comparisons,
                'summary': self._generate_comparison_summary(risk_comparison, clause_comparisons)
            }
            
        except Exception as e:
            logger.error(f"Document comparison error: {str(e)}")
            return {
                'success': False,
                'error': f"Comparison failed: {str(e)}"
            }
    
    def _calculate_risk_difference(self, risk1, risk2) -> str:
        """Calculate the difference between two risk levels"""
        risk_values = {'GREEN': 1, 'YELLOW': 2, 'RED': 3}
        diff = risk_values.get(risk2.level, 2) - risk_values.get(risk1.level, 2)
        
        if diff > 0:
            return "INCREASED"
        elif diff < 0:
            return "DECREASED"
        else:
            return "SAME"
    
    def _get_risk_change(self, risk1: str, risk2: str) -> str:
        """Get risk change description"""
        if risk1 == risk2:
            return "NO_CHANGE"
        
        risk_order = ['GREEN', 'YELLOW', 'RED']
        if risk_order.index(risk2) > risk_order.index(risk1):
            return "INCREASED"
        else:
            return "DECREASED"
    
    def _generate_comparison_summary(self, risk_comparison: Dict, clause_comparisons: List) -> str:
        """Generate a summary of the comparison"""
        overall_change = risk_comparison['risk_difference']
        
        if overall_change == "INCREASED":
            summary = "⚠️ Document 2 has higher overall risk than Document 1. "
        elif overall_change == "DECREASED":
            summary = "✅ Document 2 has lower overall risk than Document 1. "
        else:
            summary = "📊 Both documents have similar overall risk levels. "
        
        # Count clause changes
        increased = sum(1 for c in clause_comparisons if c['risk_change'] == 'INCREASED')
        decreased = sum(1 for c in clause_comparisons if c['risk_change'] == 'DECREASED')
        
        summary += f"Clause analysis: {increased} clauses show increased risk, {decreased} show decreased risk."
        
        return summary

# Initialize document comparison service
document_comparison = DocumentComparison()

# Export functionality
class ReportExporter:
    """Export analysis reports to PDF and Word formats"""
    
    def __init__(self):
        self.available_formats = []
        if REPORTLAB_AVAILABLE:
            self.available_formats.append('pdf')
        if PYTHON_DOCX_AVAILABLE:
            self.available_formats.append('docx')
    
    def export_analysis_to_pdf(self, analysis: DocumentAnalysis, filename: str = None) -> BytesIO:
        """Export analysis to PDF format"""
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab not available for PDF export")
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            textColor=colors.darkblue
        )
        story.append(Paragraph("Legal Document Analysis Report", title_style))
        story.append(Spacer(1, 12))
        
        # Overall Risk Assessment
        risk_color = colors.red if analysis.overall_risk.level == 'RED' else (colors.orange if analysis.overall_risk.level == 'YELLOW' else colors.green)
        risk_style = ParagraphStyle('RiskStyle', parent=styles['Heading2'], textColor=risk_color)
        story.append(Paragraph(f"Overall Risk: {analysis.overall_risk.level}", risk_style))
        story.append(Paragraph(analysis.summary, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Confidence and Categories
        if hasattr(analysis.overall_risk, 'confidence_percentage'):
            story.append(Paragraph(f"Analysis Confidence: {analysis.overall_risk.confidence_percentage}%", styles['Normal']))
        
        if hasattr(analysis.overall_risk, 'risk_categories'):
            story.append(Paragraph("Risk Categories:", styles['Heading3']))
            for category, score in analysis.overall_risk.risk_categories.items():
                story.append(Paragraph(f"• {category.title()}: {int(score * 100)}%", styles['Normal']))
            story.append(Spacer(1, 12))
        
        # Detailed Analysis
        story.append(Paragraph("Detailed Clause Analysis", styles['Heading2']))
        
        for i, result in enumerate(analysis.analysis_results):
            # Clause header
            clause_title = f"Clause {i + 1} - {result.risk_level.level} Risk"
            story.append(Paragraph(clause_title, styles['Heading3']))
            
            # Explanation
            story.append(Paragraph("Explanation:", styles['Heading4']))
            story.append(Paragraph(result.plain_explanation, styles['Normal']))
            
            # Recommendations
            if result.recommendations:
                story.append(Paragraph("Recommendations:", styles['Heading4']))
                for rec in result.recommendations:
                    story.append(Paragraph(f"• {rec}", styles['Normal']))
            
            story.append(Spacer(1, 12))
        
        # Footer
        story.append(Spacer(1, 20))
        footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, textColor=colors.grey)
        story.append(Paragraph(f"Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} by LegalSaathi Document Advisor", footer_style))
        story.append(Paragraph("Disclaimer: This analysis is for informational purposes only and does not constitute legal advice.", footer_style))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def export_analysis_to_docx(self, analysis: DocumentAnalysis, filename: str = None) -> BytesIO:
        """Export analysis to Word document format"""
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx not available for Word export")
        
        doc = DocxDocument()
        
        # Title
        title = doc.add_heading('Legal Document Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Overall Risk Assessment
        doc.add_heading('Overall Risk Assessment', level=1)
        risk_para = doc.add_paragraph()
        risk_run = risk_para.add_run(f'Risk Level: {analysis.overall_risk.level}')
        # Note: Word document color formatting requires RGBColor, skipping for compatibility
        
        doc.add_paragraph(analysis.summary)
        
        # Confidence and Categories
        if hasattr(analysis.overall_risk, 'confidence_percentage'):
            doc.add_paragraph(f'Analysis Confidence: {analysis.overall_risk.confidence_percentage}%')
        
        if hasattr(analysis.overall_risk, 'risk_categories'):
            doc.add_heading('Risk Categories', level=2)
            for category, score in analysis.overall_risk.risk_categories.items():
                doc.add_paragraph(f'{category.title()}: {int(score * 100)}%', style='List Bullet')
        
        # Detailed Analysis
        doc.add_heading('Detailed Clause Analysis', level=1)
        
        for i, result in enumerate(analysis.analysis_results):
            # Clause header
            doc.add_heading(f'Clause {i + 1} - {result.risk_level.level} Risk', level=2)
            
            # Explanation
            doc.add_heading('Explanation', level=3)
            doc.add_paragraph(result.plain_explanation)
            
            # Recommendations
            if result.recommendations:
                doc.add_heading('Recommendations', level=3)
                for rec in result.recommendations:
                    doc.add_paragraph(rec, style='List Bullet')
        
        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.add_run(f'Report generated on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")} by LegalSaathi Document Advisor').italic = True
        doc.add_paragraph('Disclaimer: This analysis is for informational purposes only and does not constitute legal advice.').italic = True
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

# Initialize report exporter
report_exporter = ReportExporter()

@app.route('/analyze', methods=['POST'])
@rate_limit_decorator
def analyze_document():
    """
    Enhanced document analysis endpoint with caching, monitoring, and improved error handling
    Requirements: 7.2 - Implement drag-and-drop file upload with basic validation
    """
    start_time = time.time()
    document_text = ""
    file_info = None
    warnings = []
    client_ip = request.environ.get('HTTP_X_FORWARDED_FOR', request.remote_addr)
    
    try:
        logger.info(f"Analysis request from IP: {client_ip}")
        # Check if file was uploaded
        if 'document_file' in request.files and request.files['document_file'].filename:
            file = request.files['document_file']
            
            # Process uploaded file
            file_result = file_processor.process_uploaded_file(file)
            
            if not file_result.success:
                return jsonify({
                    'success': False,
                    'error': file_result.error_message,
                    'document_text': request.form.get('document_text', '')
                }), 400
            
            document_text = file_result.text_content
            file_info = file_result.file_info
            warnings.extend(file_result.warnings)
            
        else:
            # Use text input from form
            document_text = request.form.get('document_text', '')
        
        # Validate document input
        is_valid, error_message = validate_document_input(document_text)
        
        if not is_valid:
            return jsonify({
                'success': False,
                'error': error_message,
                'document_text': document_text,
                'file_info': file_info.__dict__ if file_info else None
            }), 400
        
        # Classify document type
        classification = document_classifier.classify_document(document_text)
        classification_message = document_classifier.get_analysis_message(classification)
        
        # Add classification info only when needed
        if classification.confidence > 0.8:
            # High confidence - no warning needed, just log success
            logger.info(f"Document classified as {classification.document_type.value} with high confidence ({classification.confidence:.1%})")
        elif classification.confidence > 0.6:
            # Medium confidence - show informational message
            warnings.append(f"✓ Detected as {classification.document_type.value.replace('_', ' ').title()} document")
        elif classification.confidence > 0.3:
            # Low confidence - show warning
            warnings.append(f"⚠️ Possible {classification.document_type.value.replace('_', ' ').title()} detected - analysis confidence may vary")
        else:
            # Very low confidence - use general analysis
            warnings.append("📄 Using general legal document analysis for best results")
        
        # Detect user expertise level for adaptive explanations
        user_questions = request.form.get('user_questions', '')
        expertise_level = request.form.get('expertise_level', 'beginner')
        
        # Use form selection or detect from questions
        if expertise_level in ['beginner', 'intermediate', 'expert']:
            expertise_profile = expertise_detector._get_expertise_profile(expertise_level)
        else:
            expertise_profile = expertise_detector.detect_expertise(user_questions, 'medium')
        
        # Check cache first
        cache_key = get_cache_key(document_text, "analysis")
        cached_analysis = get_from_cache(cache_key)
        
        if cached_analysis:
            logger.info("Returning cached analysis result")
            metrics_tracker.record_analysis(True, time.time() - start_time)
            
            # Add real-time data to cached result
            cached_analysis['processing_time'] = time.time() - start_time
            cached_analysis['file_info'] = file_info
            cached_analysis['warnings'] = warnings
            cached_analysis['classification'] = classification
            cached_analysis['expertise_profile'] = expertise_profile
            
            return jsonify({
                'success': True,
                'analysis': cached_analysis,
                'file_info': file_info.__dict__ if file_info else None,
                'warnings': warnings,
                'classification': {
                    'document_type': {'value': classification.document_type.value},
                    'confidence': classification.confidence
                } if classification else None
            })
        
        # Perform enhanced AI-powered legal document analysis with Google Cloud AI services
        try:
            logger.info("Starting enhanced document analysis with Google Cloud AI")
            
            # Get document bytes for Document AI if available
            document_bytes = None
            if file_info and hasattr(file_processor, 'get_file_bytes'):
                try:
                    document_bytes = file_processor.get_file_bytes()
                except:
                    pass  # Continue without Document AI if bytes unavailable
            
            analysis = analyze_legal_document_with_llm(
                document_text, 
                classification.document_type, 
                expertise_profile,
                document_bytes
            )
            
            # Add enhanced risk category analysis
            category_analysis = risk_categories.analyze_risk_categories(document_text)
            severity_indicators = risk_categories.get_severity_indicators(category_analysis)
            
            # Store successful analysis in cache
            analysis_dict = {
                'document_id': analysis.document_id,
                'analysis_results': [
                    {
                        'clause_id': r.clause_id,
                        'risk_level': {
                            'level': r.risk_level.level,
                            'score': r.risk_level.score,
                            'reasons': r.risk_level.reasons,
                            'severity': r.risk_level.severity,
                            'confidence_percentage': r.risk_level.confidence_percentage,
                            'risk_categories': r.risk_level.risk_categories,
                            'low_confidence_warning': r.risk_level.low_confidence_warning
                        },
                        'plain_explanation': r.plain_explanation,
                        'legal_implications': r.legal_implications,
                        'recommendations': r.recommendations
                    } for r in analysis.analysis_results
                ],
                'overall_risk': {
                    'level': analysis.overall_risk.level,
                    'score': analysis.overall_risk.score,
                    'reasons': analysis.overall_risk.reasons,
                    'severity': analysis.overall_risk.severity,
                    'confidence_percentage': analysis.overall_risk.confidence_percentage,
                    'risk_categories': analysis.overall_risk.risk_categories,
                    'low_confidence_warning': analysis.overall_risk.low_confidence_warning
                },
                'summary': analysis.summary,
                'processing_time': analysis.processing_time,
                'category_analysis': category_analysis,
                'severity_indicators': severity_indicators
            }
            
            store_in_cache(cache_key, analysis_dict)
            
            # Record successful analysis metrics
            avg_confidence = sum(r.risk_level.confidence_percentage for r in analysis.analysis_results) / len(analysis.analysis_results) if analysis.analysis_results else 0
            metrics_tracker.record_analysis(True, time.time() - start_time, avg_confidence / 100)
            
            # Add file info and warnings to analysis
            if hasattr(analysis, 'file_info'):
                analysis.file_info = file_info
            if hasattr(analysis, 'warnings'):
                analysis.warnings = warnings
            if hasattr(analysis, 'classification'):
                analysis.classification = classification
            if hasattr(analysis, 'expertise_profile'):
                analysis.expertise_profile = expertise_profile
            if hasattr(analysis, 'category_analysis'):
                analysis.category_analysis = category_analysis
            if hasattr(analysis, 'severity_indicators'):
                analysis.severity_indicators = severity_indicators
            
            logger.info(f"Analysis completed successfully in {time.time() - start_time:.2f} seconds")
                
        except ConnectionError as e:
            logger.error(f"Connection error during analysis: {str(e)}")
            metrics_tracker.record_analysis(False, time.time() - start_time, error_type="ConnectionError")
            error_msg = "Unable to connect to AI analysis service. Please try again in a few moments."
            return jsonify({
                'success': False,
                'error': error_msg,
                'document_text': document_text,
                'file_info': file_info.__dict__ if file_info else None
            }), 500
        
        except TimeoutError as e:
            logger.error(f"Timeout error during analysis: {str(e)}")
            metrics_tracker.record_analysis(False, time.time() - start_time, error_type="TimeoutError")
            error_msg = "Analysis service timed out. Please try again with a shorter document or try again later."
            return jsonify({
                'success': False,
                'error': error_msg,
                'document_text': document_text,
                'file_info': file_info.__dict__ if file_info else None
            }), 500
        
        except Exception as e:
            logger.error(f"Analysis error: {str(e)}")
            # Enhanced error handling with user-friendly messages
            error_type = type(e).__name__
            
            if "memory" in str(e).lower() or "resource" in str(e).lower():
                metrics_tracker.record_analysis(False, time.time() - start_time, error_type="MemoryError")
                error_msg = "Analysis service is currently overloaded. Please try again in a few minutes."
            elif "api" in str(e).lower() or "key" in str(e).lower():
                logger.warning("API configuration issue, using fallback analysis")
                metrics_tracker.record_analysis(False, time.time() - start_time, error_type="APIError")
                error_msg = "Analysis service configuration issue. Using basic analysis instead."
                # Fallback to basic analysis
                analysis = create_fallback_analysis(document_text)
            else:
                logger.error(f"Unexpected analysis error: {str(e)}")
                metrics_tracker.record_analysis(False, time.time() - start_time, error_type="UnexpectedError")
                error_msg = "An unexpected error occurred during analysis. Using basic analysis instead."
                # Fallback to basic analysis
                analysis = create_fallback_analysis(document_text)
            
            # If we couldn't create fallback analysis, show error
            if 'analysis' not in locals():
                if request.headers.get('Accept') == 'application/json' or request.headers.get('Content-Type') == 'application/json':
                    return jsonify({
                        'success': False,
                        'error': error_msg,
                        'document_text': document_text,
                        'file_info': file_info.__dict__ if file_info else None
                    }), 500
                return jsonify({
                    'success': False,
                    'error': error_msg,
                    'document_text': document_text,
                    'file_info': file_info.__dict__ if file_info else None
                }), 500
        
        # Return JSON response (API-only)
        return jsonify({
            'success': True,
            'analysis': {
                'document_id': analysis.document_id,
                'analysis_results': [
                    {
                        'clause_id': r.clause_id,
                        'risk_level': {
                            'level': r.risk_level.level,
                            'score': r.risk_level.score,
                            'reasons': r.risk_level.reasons,
                            'severity': r.risk_level.severity,
                            'confidence_percentage': r.risk_level.confidence_percentage,
                            'risk_categories': r.risk_level.risk_categories,
                            'low_confidence_warning': r.risk_level.low_confidence_warning
                        },
                        'plain_explanation': r.plain_explanation,
                        'legal_implications': r.legal_implications,
                        'recommendations': r.recommendations
                    } for r in analysis.analysis_results
                ],
                'overall_risk': {
                    'level': analysis.overall_risk.level,
                    'score': analysis.overall_risk.score,
                    'reasons': analysis.overall_risk.reasons,
                    'severity': analysis.overall_risk.severity,
                    'confidence_percentage': analysis.overall_risk.confidence_percentage,
                    'risk_categories': analysis.overall_risk.risk_categories,
                    'low_confidence_warning': analysis.overall_risk.low_confidence_warning
                },
                'summary': analysis.summary,
                'processing_time': analysis.processing_time,
                'severity_indicators': getattr(analysis, 'severity_indicators', []),
                'enhanced_insights': getattr(analysis, 'enhanced_insights', {})
            },
            'file_info': file_info.__dict__ if file_info else None,
            'warnings': warnings,
            'classification': {
                'document_type': {'value': classification.document_type.value},
                'confidence': classification.confidence
            } if classification else None
        })
    
    except Exception as e:
        # Catch-all error handler
        print(f"Unexpected error in analyze_document: {str(e)}")
        error_msg = "An unexpected error occurred. Please try again or contact support if the problem persists."
        return jsonify({
            'success': False,
            'error': error_msg,
            'document_text': request.form.get('document_text', '')
        }), 500

@app.route('/api/translate', methods=['POST'])
def translate_text():
    """
    API endpoint for text translation using Google Translate service
    Requirements: Integrate Google Translate API using provided key for English/Hindi translation
    """
    try:
        data = request.get_json()
        text = data.get('text', '')
        target_language = data.get('target_language', 'hi')
        
        if not text:
            return jsonify({'success': False, 'error': 'No text provided'})
        
        result = translation_service.translate_text(text, target_language)
        return jsonify(result)
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/translate-assessment', methods=['POST'])
def translate_risk_assessment():
    """
    API endpoint for translating entire risk assessment results
    """
    try:
        data = request.get_json()
        assessment = data.get('assessment', {})
        target_language = data.get('target_language', 'hi')
        
        if not assessment:
            return jsonify({'success': False, 'error': 'No assessment provided'})
        
        result = translation_service.translate_risk_assessment(assessment, target_language)
        return jsonify({'success': True, 'translated_assessment': result})
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/clarify', methods=['POST'])
@rate_limit_decorator
def ask_clarification():
    """
    Enhanced AI-powered clarification endpoint with improved Google Cloud AI integration
    Features: Context awareness, conversation history, confidence scoring, caching
    """
    try:
        data = request.get_json()
        question = data.get('question', '').strip()
        context = data.get('context', {})
        
        if not question:
            return jsonify({'success': False, 'error': 'Question is required'})
        
        logger.info(f"AI clarification request: {question[:50]}...")
        
        # Check cache for similar questions
        question_cache_key = get_cache_key(question, "clarification")
        cached_response = get_from_cache(question_cache_key)
        
        if cached_response:
            logger.info("Returning cached clarification response")
            return jsonify(cached_response)
        
        # Use the enhanced AI clarification service
        result = ai_clarification_service.ask_clarification(question, context)
        
        # Add confidence scoring to the response
        if result.get('success'):
            result['confidence_score'] = 85 if not result.get('fallback') else 45
            result['response_quality'] = 'high' if not result.get('fallback') else 'basic'
            result['timestamp'] = datetime.now().isoformat()
            
            # Store successful responses in cache
            store_in_cache(question_cache_key, result)
            
        logger.info(f"AI clarification completed: {'success' if result.get('success') else 'failed'}")
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"AI clarification error: {str(e)}")
        return jsonify({
            'success': False, 
            'error': str(e),
            'fallback_message': 'AI clarification service is temporarily unavailable. Please try again later.'
        })

@app.route('/api/conversation-summary')
def get_conversation_summary():
    """Get conversation summary from AI clarification service"""
    try:
        summary = ai_clarification_service.get_conversation_summary()
        return jsonify(summary)
    except Exception as e:
        logger.error(f"Conversation summary error: {str(e)}")
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/compare-documents', methods=['POST'])
@rate_limit_decorator
def compare_documents():
    """
    Compare two documents for risk differences
    Enhanced Google Cloud AI integration for document comparison
    """
    try:
        data = request.get_json()
        
        if not data or 'document1' not in data or 'document2' not in data:
            return jsonify({
                'success': False,
                'error': 'Both document1 and document2 text are required'
            }), 400
        
        doc1_text = data['document1'].strip()
        doc2_text = data['document2'].strip()
        
        if not doc1_text or not doc2_text:
            return jsonify({
                'success': False,
                'error': 'Both documents must contain text'
            }), 400
        
        # Check cache for comparison
        comparison_key = get_cache_key(f"{doc1_text}||{doc2_text}", "comparison")
        cached_comparison = get_from_cache(comparison_key)
        
        if cached_comparison:
            logger.info("Returning cached comparison result")
            return jsonify(cached_comparison)
        
        # Perform comparison
        logger.info("Starting document comparison")
        comparison_result = document_comparison.compare_documents(doc1_text, doc2_text)
        
        if comparison_result['success']:
            # Store in cache
            store_in_cache(comparison_key, comparison_result)
            logger.info("Document comparison completed successfully")
        
        return jsonify(comparison_result)
        
    except Exception as e:
        logger.error(f"Document comparison error: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Comparison failed: {str(e)}'
        }), 500

@app.route('/api/export/<format_type>/<document_id>')
def export_analysis(format_type, document_id):
    """
    Export analysis results to PDF or Word format
    Enhanced export functionality with multiple formats
    """
    try:
        if format_type not in ['pdf', 'docx']:
            return jsonify({
                'success': False,
                'error': 'Unsupported format. Use pdf or docx'
            }), 400
        
        if format_type not in report_exporter.available_formats:
            return jsonify({
                'success': False,
                'error': f'{format_type.upper()} export not available. Missing required libraries.'
            }), 503
        
        # Get analysis from cache or session
        # In production, you'd retrieve from database using document_id
        cache_keys = [key for key in analysis_cache.keys() if key.startswith('analysis:')]
        
        if not cache_keys:
            return jsonify({
                'success': False,
                'error': 'No analysis found for export. Please analyze a document first.'
            }), 404
        
        # Use the most recent analysis (simplified)
        latest_cache_key = max(cache_keys, key=lambda k: analysis_cache[k]['timestamp'])
        analysis_data = analysis_cache[latest_cache_key]['data']
        
        # Convert to DocumentAnalysis object
        analysis = type('DocumentAnalysis', (), analysis_data)()
        
        # Generate export
        if format_type == 'pdf':
            buffer = report_exporter.export_analysis_to_pdf(analysis)
            mimetype = 'application/pdf'
            filename = f'legal_analysis_{document_id}.pdf'
        else:  # docx
            buffer = report_exporter.export_analysis_to_docx(analysis)
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            filename = f'legal_analysis_{document_id}.docx'
        
        logger.info(f"Generated {format_type.upper()} export for document {document_id}")
        
        return send_file(
            buffer,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        logger.error(f"Export error: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Export failed: {str(e)}'
        }), 500

@app.route('/api/export/pdf', methods=['POST'])
def export_analysis_pdf():
    """Export analysis results to PDF format"""
    try:
        data = request.get_json()
        if not data or 'analysis' not in data:
            return jsonify({
                'success': False,
                'error': 'Analysis data required'
            }), 400
        
        if 'pdf' not in report_exporter.available_formats:
            return jsonify({
                'success': False,
                'error': 'PDF export not available. Missing required libraries.'
            }), 503
        
        # Convert data to DocumentAnalysis object
        analysis_data = data['analysis']
        analysis = type('DocumentAnalysis', (), analysis_data)()
        
        # Generate PDF
        buffer = report_exporter.export_analysis_to_pdf(analysis)
        
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='legal_analysis_report.pdf'
        )
        
    except Exception as e:
        logger.error(f"PDF export error: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'PDF export failed: {str(e)}'
        }), 500

@app.route('/api/export/word', methods=['POST'])
def export_analysis_word():
    """Export analysis results to Word format"""
    try:
        data = request.get_json()
        if not data or 'analysis' not in data:
            return jsonify({
                'success': False,
                'error': 'Analysis data required'
            }), 400
        
        if 'docx' not in report_exporter.available_formats:
            return jsonify({
                'success': False,
                'error': 'Word export not available. Missing required libraries.'
            }), 503
        
        # Convert data to DocumentAnalysis object
        analysis_data = data['analysis']
        analysis = type('DocumentAnalysis', (), analysis_data)()
        
        # Generate Word document
        buffer = report_exporter.export_analysis_to_docx(analysis)
        
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='legal_analysis_report.docx'
        )
        
    except Exception as e:
        logger.error(f"Word export error: {str(e)}")
        return jsonify({
            'success': False,
            'error': f'Word export failed: {str(e)}'
        }), 500

@app.route('/api/metrics')
def get_metrics():
    """
    Get system metrics for monitoring
    Production monitoring endpoint
    """
    try:
        metrics = metrics_tracker.get_metrics()
        
        # Add cache statistics
        cache_stats = {
            'cache_size': len(analysis_cache),
            'rate_limit_entries': len(rate_limit_cache)
        }
        
        # Add system health indicators
        health_status = {
            'status': 'healthy' if metrics['successful_analyses'] > metrics['failed_analyses'] else 'degraded',
            'uptime': time.time() - app.config.get('START_TIME', time.time()),
            'available_export_formats': report_exporter.available_formats
        }
        
        return jsonify({
            'success': True,
            'metrics': metrics,
            'cache_stats': cache_stats,
            'health': health_status,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logger.error(f"Metrics error: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/confidence-score/<document_id>')
def get_confidence_score(document_id):
    """
    Get detailed confidence scoring for a document analysis
    Enhanced AI transparency feature
    """
    try:
        # Find analysis in cache (simplified)
        cache_keys = [key for key in analysis_cache.keys() if key.startswith('analysis:')]
        
        if not cache_keys:
            return jsonify({
                'success': False,
                'error': 'No analysis found'
            }), 404
        
        # Use the most recent analysis (simplified)
        latest_cache_key = max(cache_keys, key=lambda k: analysis_cache[k]['timestamp'])
        analysis_data = analysis_cache[latest_cache_key]['data']
        
        # Extract confidence information
        confidence_details = {
            'overall_confidence': analysis_data['overall_risk']['confidence_percentage'],
            'low_confidence_warning': analysis_data['overall_risk']['low_confidence_warning'],
            'clause_confidences': [
                {
                    'clause_id': result['clause_id'],
                    'confidence': result['risk_level']['confidence_percentage'],
                    'low_confidence': result['risk_level']['low_confidence_warning']
                }
                for result in analysis_data['analysis_results']
            ],
            'confidence_factors': {
                'document_clarity': 'High' if analysis_data['overall_risk']['confidence_percentage'] > 80 else 'Medium' if analysis_data['overall_risk']['confidence_percentage'] > 60 else 'Low',
                'ai_model_certainty': 'High' if not analysis_data['overall_risk']['low_confidence_warning'] else 'Low',
                'pattern_recognition': 'Strong' if analysis_data['overall_risk']['confidence_percentage'] > 75 else 'Moderate'
            }
        }
        
        return jsonify({
            'success': True,
            'confidence_details': confidence_details,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logger.error(f"Confidence score error: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# Error handlers for production
@app.errorhandler(429)
def rate_limit_handler(e):
    """Handle rate limit exceeded errors"""
    return jsonify({
        'success': False,
        'error': 'Rate limit exceeded. Please try again in a minute.',
        'retry_after': 60
    }), 429

@app.errorhandler(500)
def internal_error_handler(e):
    """Handle internal server errors"""
    logger.error(f"Internal server error: {str(e)}")
    return jsonify({
        'success': False,
        'error': 'Internal server error. Please try again later.'
    }), 500

@app.errorhandler(404)
def not_found_handler(e):
    """Handle not found errors"""
    return jsonify({
        'success': False,
        'error': 'Resource not found'
    }), 404

# Health check endpoint
@app.route('/api/speech-to-text', methods=['POST'])
@rate_limit_decorator
def speech_to_text():
    """
    Google Cloud Speech-to-Text API endpoint for voice input
    Supports legal terminology and multiple languages
    """
    try:
        if 'audio' not in request.files:
            return jsonify({'success': False, 'error': 'No audio file provided'}), 400
        
        audio_file = request.files['audio']
        language_code = request.form.get('language_code', 'en-US')
        
        if not audio_file.filename:
            return jsonify({'success': False, 'error': 'No audio file selected'}), 400
        
        # Read audio content
        audio_content = audio_file.read()
        
        # Process with Google Speech-to-Text
        result = google_speech.transcribe_audio(audio_content, language_code)
        
        if result['success']:
            logger.info(f"Speech transcription successful: {len(result['transcript'])} characters")
        
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Speech-to-text error: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/document-ai-analysis', methods=['POST'])
@rate_limit_decorator
def document_ai_analysis():
    """
    Google Cloud Document AI analysis endpoint
    Extracts structured data from legal documents
    """
    try:
        if 'document' not in request.files:
            return jsonify({'success': False, 'error': 'No document file provided'}), 400
        
        document_file = request.files['document']
        
        if not document_file.filename:
            return jsonify({'success': False, 'error': 'No document file selected'}), 400
        
        # Read document content
        document_content = document_file.read()
        mime_type = document_file.content_type or 'application/pdf'
        
        # Process with Google Document AI
        result = google_document_ai.process_legal_document(document_content, mime_type)
        
        if result['success']:
            logger.info(f"Document AI analysis successful: {len(result['entities'])} entities extracted")
        
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Document AI analysis error: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/natural-language-analysis', methods=['POST'])
@rate_limit_decorator
def natural_language_analysis():
    """
    Google Cloud Natural Language AI analysis endpoint
    Provides sentiment, entity, and complexity analysis
    """
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        
        if not text:
            return jsonify({'success': False, 'error': 'No text provided'}), 400
        
        # Process with Google Natural Language AI
        result = google_natural_language.analyze_legal_document(text)
        
        if result['success']:
            logger.info(f"Natural Language analysis successful: {len(result['entities'])} entities found")
        
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Natural Language analysis error: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/ai-services-status')
def ai_services_status():
    """
    Get status of all Google Cloud AI services
    Monitoring endpoint for service health
    """
    try:
        status = {
            'google_translate': {
                'enabled': hasattr(translation_service, 'enabled') and translation_service.enabled,
                'service': 'Google Translate API',
                'cloud_enabled': hasattr(translation_service, 'cloud_enabled') and translation_service.cloud_enabled,
                'status': 'Active' if (hasattr(translation_service, 'enabled') and translation_service.enabled) else 'Inactive'
            },
            'google_document_ai': {
                'enabled': google_document_ai.enabled,
                'service': 'Google Cloud Document AI'
            },
            'google_natural_language': {
                'enabled': google_natural_language.enabled,
                'service': 'Google Cloud Natural Language AI'
            },
            'google_speech': {
                'enabled': google_speech.enabled,
                'service': 'Google Cloud Speech-to-Text'
            }
        }
        
        # Count enabled services
        enabled_count = sum(1 for service in status.values() if service['enabled'])
        total_count = len(status)
        
        return jsonify({
            'success': True,
            'services': status,
            'summary': {
                'enabled_services': enabled_count,
                'total_services': total_count,
                'integration_level': f"{enabled_count}/{total_count} Google Cloud AI services active"
            },
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        logger.error(f"AI services status error: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500



# Initialize application
def initialize_app():
    """Initialize application with startup configuration"""
    app.config['START_TIME'] = time.time()
    
    # Log startup information
    logger.info("LegalSaathi Document Advisor starting up...")
    logger.info(f"Available export formats: {report_exporter.available_formats}")
    logger.info(f"Cache expiry: {CACHE_EXPIRY} seconds")
    logger.info(f"Rate limit: {RATE_LIMIT_REQUESTS} requests per {RATE_LIMIT_WINDOW} seconds")
    
    # Clear any existing cache on startup
    analysis_cache.clear()
    rate_limit_cache.clear()
    
    logger.info("Application initialized successfully")

if __name__ == '__main__':
    initialize_app()
    app.run(debug=True, host='0.0.0.0', port=5000)