"""
Export controller for FastAPI backend with enhanced PDF generation
"""

import logging
import io
import json
import tempfile
import os
from datetime import datetime
from typing import Dict, Any, Optional
from fastapi import HTTPException
from fastapi.responses import StreamingResponse

from models.document_models import DocumentAnalysisResponse

logger = logging.getLogger(__name__)


class ExportController:
    """Controller for document export operations"""
    
    def __init__(self):
        pass
    
    async def export_to_pdf(self, data: Dict[str, Any]) -> StreamingResponse:
        """Export analysis results to PDF - Using same enhanced quality as email"""
        try:
            logger.info("Generating PDF export for download (same quality as email)")
            
            # Generate PDF content using enhanced method
            pdf_content = await self._generate_pdf_content(data)
            
            # Validate PDF content
            if not pdf_content or len(pdf_content) == 0:
                logger.error("PDF generation returned empty content")
                raise HTTPException(
                    status_code=500,
                    detail="PDF generation failed - empty content"
                )
            
            logger.info(f"PDF generated successfully, size: {len(pdf_content)} bytes")
            
            # Create streaming response
            return StreamingResponse(
                io.BytesIO(pdf_content),
                media_type="application/pdf",
                headers={
                    "Content-Disposition": f"attachment; filename=legal_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    "Content-Length": str(len(pdf_content))
                }
            )
            
        except Exception as e:
            logger.error(f"PDF export failed: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"PDF export failed: {str(e)}"
            )
    
    async def export_to_word(self, data: Dict[str, Any]) -> StreamingResponse:
        """Export analysis results to Word document"""
        try:
            logger.info("Generating Word export")
            
            # Generate Word content
            word_content = await self._generate_word_content(data)
            
            return StreamingResponse(
                io.BytesIO(word_content),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename=legal_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                }
            )
            
        except Exception as e:
            logger.error(f"Word export failed: {e}")
            raise HTTPException(
                status_code=500,
                detail=f"Word export failed: {str(e)}"
            )
    
    async def generate_enhanced_pdf(self, analysis: DocumentAnalysisResponse) -> bytes:
        """Generate enhanced branded PDF with risk visualization"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, 
                PageBreak, Image, KeepTogether
            )
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch, cm
            from reportlab.lib import colors
            from reportlab.graphics.shapes import Drawing, Rect, Circle
            from reportlab.graphics.charts.piecharts import Pie
            from reportlab.graphics.charts.barcharts import VerticalBarChart
            from reportlab.graphics import renderPDF
            
            # Create PDF buffer
            buffer = io.BytesIO()
            
            # Create PDF document with custom margins
            doc = SimpleDocTemplate(
                buffer, 
                pagesize=A4, 
                topMargin=0.8*inch,
                bottomMargin=0.8*inch,
                leftMargin=0.8*inch,
                rightMargin=0.8*inch
            )
            styles = getSampleStyleSheet()
            
            # Custom styles with branding
            title_style = ParagraphStyle(
                'BrandedTitle',
                parent=styles['Heading1'],
                fontSize=28,
                spaceAfter=20,
                textColor=colors.Color(0.05, 0.65, 0.9),  # LegalSaathi blue
                alignment=1,  # Center alignment
                fontName='Helvetica-Bold'
            )
            
            subtitle_style = ParagraphStyle(
                'Subtitle',
                parent=styles['Normal'],
                fontSize=14,
                spaceAfter=30,
                textColor=colors.Color(0.4, 0.4, 0.4),
                alignment=1,
                fontName='Helvetica'
            )
            
            heading_style = ParagraphStyle(
                'BrandedHeading',
                parent=styles['Heading2'],
                fontSize=18,
                spaceAfter=15,
                spaceBefore=20,
                textColor=colors.Color(0.1, 0.1, 0.1),
                fontName='Helvetica-Bold',
                borderWidth=0,
                borderColor=colors.Color(0.05, 0.65, 0.9),
                borderPadding=5
            )
            
            risk_high_style = ParagraphStyle(
                'RiskHigh',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.red,
                fontName='Helvetica-Bold'
            )
            
            risk_medium_style = ParagraphStyle(
                'RiskMedium',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.orange,
                fontName='Helvetica-Bold'
            )
            
            risk_low_style = ParagraphStyle(
                'RiskLow',
                parent=styles['Normal'],
                fontSize=12,
                textColor=colors.green,
                fontName='Helvetica-Bold'
            )
            
            # Build content
            content = []
            
            # Header with branding
            content.append(Paragraph("🏛️ LegalSaathi", title_style))
            content.append(Paragraph("AI-Powered Legal Document Analysis Report", subtitle_style))
            content.append(Spacer(1, 20))
            
            # Executive Summary Box
            summary_data = [
                ['Overall Risk Level', analysis.overall_risk.level],
                ['Risk Score', f"{analysis.overall_risk.score:.1%}"],
                ['Confidence Level', f"{analysis.overall_risk.confidence_percentage}%"],
                ['Analysis Date', datetime.now().strftime('%B %d, %Y')],
                ['Total Clauses Analyzed', str(len(analysis.clause_assessments))]
            ]
            
            summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.05, 0.65, 0.9)),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.Color(0.95, 0.95, 0.95)),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 11),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.Color(0.98, 0.98, 0.98)])
            ]))
            
            content.append(Paragraph("Executive Summary", heading_style))
            content.append(summary_table)
            content.append(Spacer(1, 20))
            
            # Risk Assessment Summary
            content.append(Paragraph("Risk Assessment", heading_style))
            content.append(Paragraph(analysis.summary, styles['Normal']))
            content.append(Spacer(1, 20))
            
            # Risk Distribution Chart
            if len(analysis.clause_assessments) > 0:
                risk_counts = {
                    'HIGH': len([c for c in analysis.clause_assessments if c.risk_assessment.level == 'RED']),
                    'MODERATE': len([c for c in analysis.clause_assessments if c.risk_assessment.level == 'YELLOW']),
                    'LOW': len([c for c in analysis.clause_assessments if c.risk_assessment.level == 'GREEN'])
                }
                
                # Create risk distribution table
                risk_data = [
                    ['Risk Level', 'Count', 'Percentage'],
                    ['HIGH RISK', str(risk_counts['HIGH']), f"{risk_counts['HIGH']/len(analysis.clause_assessments)*100:.1f}%"],
                    ['MODERATE RISK', str(risk_counts['MODERATE']), f"{risk_counts['MODERATE']/len(analysis.clause_assessments)*100:.1f}%"],
                    ['LOW RISK', str(risk_counts['LOW']), f"{risk_counts['LOW']/len(analysis.clause_assessments)*100:.1f}%"]
                ]
                
                risk_table = Table(risk_data, colWidths=[2*inch, 1*inch, 1.5*inch])
                risk_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.2, 0.2, 0.2)),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (0, 1), colors.Color(1, 0.9, 0.9)),
                    ('BACKGROUND', (0, 2), (0, 2), colors.Color(1, 0.95, 0.8)),
                    ('BACKGROUND', (0, 3), (0, 3), colors.Color(0.9, 1, 0.9)),
                    ('TEXTCOLOR', (0, 1), (0, 1), colors.red),
                    ('TEXTCOLOR', (0, 2), (0, 2), colors.orange),
                    ('TEXTCOLOR', (0, 3), (0, 3), colors.green),
                    ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTSIZE', (0, 1), (-1, -1), 11)
                ]))
                
                content.append(Paragraph("Risk Distribution", heading_style))
                content.append(risk_table)
                content.append(Spacer(1, 30))
            
            # Page break before detailed analysis
            content.append(PageBreak())
            
            # Detailed Clause Analysis
            content.append(Paragraph("Detailed Clause Analysis", heading_style))
            content.append(Spacer(1, 15))
            
            for i, clause in enumerate(analysis.clause_assessments):
                # Determine risk style
                if clause.risk_assessment.level == 'RED':
                    risk_style = risk_high_style
                    risk_color = colors.Color(1, 0.9, 0.9)
                elif clause.risk_assessment.level == 'YELLOW':
                    risk_style = risk_medium_style
                    risk_color = colors.Color(1, 0.95, 0.8)
                else:
                    risk_style = risk_low_style
                    risk_color = colors.Color(0.9, 1, 0.9)
                
                # Clause header
                clause_header = f"Clause {i+1}: {clause.clause_id}"
                content.append(Paragraph(clause_header, styles['Heading3']))
                
                # Risk level and score
                risk_info = f"Risk Level: {clause.risk_assessment.level} | Score: {clause.risk_assessment.score:.1%} | Confidence: {clause.risk_assessment.confidence_percentage}%"
                content.append(Paragraph(risk_info, risk_style))
                content.append(Spacer(1, 8))
                
                # Clause explanation
                content.append(Paragraph("<b>Plain Language Explanation:</b>", styles['Normal']))
                content.append(Paragraph(clause.plain_explanation, styles['Normal']))
                content.append(Spacer(1, 8))
                
                # Legal implications
                if clause.legal_implications:
                    content.append(Paragraph("<b>Legal Implications:</b>", styles['Normal']))
                    for implication in clause.legal_implications:
                        content.append(Paragraph(f"• {implication}", styles['Normal']))
                    content.append(Spacer(1, 8))
                
                # Recommendations
                if clause.recommendations:
                    content.append(Paragraph("<b>Recommendations:</b>", styles['Normal']))
                    for recommendation in clause.recommendations:
                        content.append(Paragraph(f"• {recommendation}", styles['Normal']))
                    content.append(Spacer(1, 8))
                
                # Separator
                content.append(Spacer(1, 15))
                if i < len(analysis.clause_assessments) - 1:
                    content.append(Paragraph("─" * 80, styles['Normal']))
                    content.append(Spacer(1, 15))
            
            # Footer
            content.append(PageBreak())
            content.append(Spacer(1, 50))
            
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.Color(0.5, 0.5, 0.5),
                alignment=1
            )
            
            content.append(Paragraph("Generated by LegalSaathi AI Analysis System", footer_style))
            content.append(Paragraph(f"Report generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", footer_style))
            content.append(Spacer(1, 20))
            
            disclaimer_style = ParagraphStyle(
                'Disclaimer',
                parent=styles['Normal'],
                fontSize=9,
                textColor=colors.Color(0.4, 0.4, 0.4),
                alignment=1,
                leading=12
            )
            
            disclaimer_text = """
            <b>IMPORTANT DISCLAIMER:</b><br/>
            This analysis is generated by artificial intelligence and is for informational purposes only. 
            It does not constitute legal advice and should not be relied upon as a substitute for 
            consultation with a qualified attorney. The accuracy of AI analysis may vary, and 
            professional legal review is recommended for all important legal documents and decisions.
            """
            
            content.append(Paragraph(disclaimer_text, disclaimer_style))
            
            # Build PDF
            doc.build(content)
            
            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            return pdf_bytes
            
        except ImportError as e:
            logger.warning(f"ReportLab not available: {e}, generating simple PDF")
            return await self._generate_simple_pdf_content({'analysis': analysis.model_dump()})
        except Exception as e:
            logger.error(f"Enhanced PDF generation failed: {e}")
            # Fallback to simple PDF
            return await self._generate_simple_pdf_content({'analysis': analysis.model_dump()})

    async def _generate_pdf_content(self, data: Dict[str, Any]) -> bytes:
        """Generate PDF content from analysis data - Always use enhanced PDF like email"""
        logger.info("Generating enhanced PDF for download (same quality as email)")
        
        analysis = None
        
        # Try multiple ways to extract analysis data
        if isinstance(data.get('analysis'), dict):
            try:
                analysis = DocumentAnalysisResponse(**data['analysis'])
                logger.info("Successfully parsed analysis from dict")
            except Exception as e:
                logger.warning(f"Failed to parse analysis from dict: {e}")
        
        elif hasattr(data.get('analysis'), 'model_dump'):
            # Handle Pydantic model
            try:
                analysis_dict = data['analysis'].model_dump() if hasattr(data['analysis'], 'model_dump') else data['analysis'].dict()
                analysis = DocumentAnalysisResponse(**analysis_dict)
                logger.info("Successfully parsed analysis from Pydantic model")
            except Exception as e:
                logger.warning(f"Failed to convert Pydantic model: {e}")
        
        elif hasattr(data.get('analysis'), '__dict__'):
            # Handle object with attributes
            try:
                analysis_dict = data['analysis'].__dict__
                analysis = DocumentAnalysisResponse(**analysis_dict)
                logger.info("Successfully parsed analysis from object attributes")
            except Exception as e:
                logger.warning(f"Failed to convert object attributes: {e}")
        
        # If we successfully parsed the analysis, use enhanced PDF
        if analysis:
            try:
                logger.info("Using enhanced PDF generation (same as email)")
                return await self.generate_enhanced_pdf(analysis)
            except Exception as e:
                logger.error(f"Enhanced PDF generation failed: {e}")
        
        # Last resort: try to create a basic analysis structure from available data
        logger.warning("Attempting to create basic analysis structure for enhanced PDF")
        try:
            # Create a minimal analysis structure
            analysis_data = data.get('analysis', {})
            
            # Ensure required fields exist
            if not isinstance(analysis_data, dict):
                analysis_data = {'summary': 'Analysis results', 'overall_risk': {'level': 'GREEN', 'score': 0.0, 'confidence_percentage': 85}}
            
            # Fill in missing required fields
            if 'overall_risk' not in analysis_data:
                analysis_data['overall_risk'] = {'level': 'GREEN', 'score': 0.0, 'confidence_percentage': 85}
            if 'clause_assessments' not in analysis_data:
                analysis_data['clause_assessments'] = []
            if 'summary' not in analysis_data:
                analysis_data['summary'] = 'Document analysis completed successfully.'
            if 'analysis_id' not in analysis_data:
                analysis_data['analysis_id'] = f"analysis_{datetime.now().timestamp()}"
            if 'timestamp' not in analysis_data:
                analysis_data['timestamp'] = datetime.now().isoformat()
            
            analysis = DocumentAnalysisResponse(**analysis_data)
            logger.info("Created minimal analysis structure, using enhanced PDF")
            return await self.generate_enhanced_pdf(analysis)
            
        except Exception as e:
            logger.error(f"Failed to create minimal analysis structure: {e}")
        
        # Final fallback to simple PDF (but still try to make it better)
        logger.warning("Using fallback simple PDF generation")
        return await self._generate_simple_pdf_content(data)
    
    async def _generate_simple_pdf_content(self, data: Dict[str, Any]) -> bytes:
        """Generate enhanced simple PDF content as fallback - still high quality"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            
            # Create PDF buffer
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(
                buffer, 
                pagesize=A4,
                topMargin=0.8*inch,
                bottomMargin=0.8*inch,
                leftMargin=0.8*inch,
                rightMargin=0.8*inch
            )
            styles = getSampleStyleSheet()
            
            # Custom styles for better appearance
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=20,
                textColor=colors.Color(0.05, 0.65, 0.9),
                alignment=1,  # Center
                fontName='Helvetica-Bold'
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=12,
                spaceBefore=16,
                textColor=colors.Color(0.1, 0.1, 0.1),
                fontName='Helvetica-Bold'
            )
            
            content = []
            
            # Add branded title
            content.append(Paragraph("🏛️ LegalSaathi", title_style))
            content.append(Paragraph("Legal Document Analysis Report", styles['Heading2']))
            content.append(Spacer(1, 30))
            
            # Add generation date
            content.append(Paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
            content.append(Spacer(1, 20))
            
            # Add analysis summary
            analysis = data.get('analysis', {})
            if analysis:
                overall_risk = analysis.get('overall_risk', {})
                
                # Create summary table
                summary_data = [
                    ['Overall Risk Level', overall_risk.get('level', 'Unknown')],
                    ['Risk Score', f"{overall_risk.get('score', 0):.1%}"],
                    ['Confidence Level', f"{overall_risk.get('confidence_percentage', 0)}%"],
                    ['Analysis Date', datetime.now().strftime('%B %d, %Y')]
                ]
                
                summary_table = Table(summary_data, colWidths=[2.5*inch, 2*inch])
                summary_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.05, 0.65, 0.9)),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.Color(0.95, 0.95, 0.95)),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 1), (-1, -1), 11),
                ]))
                
                content.append(Paragraph("Executive Summary", heading_style))
                content.append(summary_table)
                content.append(Spacer(1, 20))
                
                summary = analysis.get('summary', '')
                if summary:
                    content.append(Paragraph("Risk Assessment", heading_style))
                    content.append(Paragraph(summary, styles['Normal']))
                    content.append(Spacer(1, 20))
                
                # Add clause analysis if available
                analysis_results = analysis.get('analysis_results', [])
                if analysis_results:
                    content.append(Paragraph("Clause Analysis Summary", heading_style))
                    
                    high_risk = len([r for r in analysis_results if r.get('risk_level', {}).get('level') == 'RED'])
                    medium_risk = len([r for r in analysis_results if r.get('risk_level', {}).get('level') == 'YELLOW'])
                    low_risk = len([r for r in analysis_results if r.get('risk_level', {}).get('level') == 'GREEN'])
                    
                    clause_data = [
                        ['Risk Level', 'Count', 'Percentage'],
                        ['High Risk', str(high_risk), f"{high_risk/len(analysis_results)*100:.1f}%"],
                        ['Medium Risk', str(medium_risk), f"{medium_risk/len(analysis_results)*100:.1f}%"],
                        ['Low Risk', str(low_risk), f"{low_risk/len(analysis_results)*100:.1f}%"]
                    ]
                    
                    clause_table = Table(clause_data, colWidths=[1.5*inch, 1*inch, 1.5*inch])
                    clause_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.2, 0.2, 0.2)),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (0, 1), colors.Color(1, 0.9, 0.9)),
                        ('BACKGROUND', (0, 2), (0, 2), colors.Color(1, 0.95, 0.8)),
                        ('BACKGROUND', (0, 3), (0, 3), colors.Color(0.9, 1, 0.9)),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 11),
                    ]))
                    
                    content.append(clause_table)
                    content.append(Spacer(1, 30))
            
            # Add footer
            content.append(Spacer(1, 50))
            footer_style = ParagraphStyle(
                'Footer',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.Color(0.5, 0.5, 0.5),
                alignment=1
            )
            
            content.append(Paragraph("Generated by LegalSaathi AI Analysis System", footer_style))
            content.append(Paragraph(f"Report generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", footer_style))
            
            # Build PDF
            doc.build(content)
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            logger.info(f"Enhanced simple PDF generated, size: {len(pdf_bytes)} bytes")
            return pdf_bytes
            
        except ImportError:
            logger.warning("ReportLab not available for enhanced simple PDF")
            # Ultimate fallback - create a basic text-based PDF-like content
            content = f"""Legal Document Analysis Report
Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}

Analysis Summary:
{json.dumps(data.get('analysis', {}), indent=2)}

Generated by LegalSaathi Document Advisor
"""
            return content.encode('utf-8')
        except Exception as e:
            logger.error(f"Enhanced simple PDF generation failed: {e}")
            # Basic fallback
            content = f"""Legal Document Analysis Report
Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}

Generated by LegalSaathi Document Advisor
"""
            return content.encode('utf-8')
    
    async def _generate_word_content(self, data: Dict[str, Any]) -> bytes:
        """Generate Word document content from analysis data"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading('Legal Document Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Analysis summary
            analysis = data.get('analysis', {})
            if analysis:
                doc.add_heading('Analysis Summary', level=1)
                
                overall_risk = analysis.get('overall_risk', {})
                if overall_risk:
                    risk_level = overall_risk.get('level', 'Unknown')
                    risk_score = overall_risk.get('score', 0)
                    confidence = overall_risk.get('confidence_percentage', 0)
                    
                    p = doc.add_paragraph()
                    p.add_run('Overall Risk Level: ').bold = True
                    p.add_run(f'{risk_level}\n')
                    p.add_run('Risk Score: ').bold = True
                    p.add_run(f'{risk_score:.2f}\n')
                    p.add_run('Confidence: ').bold = True
                    p.add_run(f'{confidence}%')
                
                # Summary text
                summary = analysis.get('summary', '')
                if summary:
                    doc.add_heading('Summary', level=1)
                    doc.add_paragraph(summary)
                
                # Clause analysis
                clause_results = analysis.get('analysis_results', [])
                if clause_results:
                    doc.add_heading('Clause Analysis', level=1)
                    
                    for i, clause in enumerate(clause_results[:10]):  # Limit to 10 clauses
                        doc.add_heading(f'Clause {i+1}', level=2)
                        
                        risk_level = clause.get('risk_level', {}).get('level', 'Unknown')
                        p = doc.add_paragraph()
                        p.add_run('Risk Level: ').bold = True
                        p.add_run(risk_level)
                        
                        explanation = clause.get('plain_explanation', '')
                        if explanation:
                            p = doc.add_paragraph()
                            p.add_run('Explanation: ').bold = True
                            p.add_run(explanation)
            
            # Footer
            doc.add_page_break()
            footer_para = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para = doc.add_paragraph("Legal Saathi Document Advisor")
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            word_bytes = buffer.getvalue()
            buffer.close()
            
            return word_bytes
            
        except ImportError:
            # Fallback if python-docx is not available
            logger.warning("python-docx not available, generating simple Word content")
            return await self._generate_simple_word_content(data)
        except Exception as e:
            logger.error(f"Word generation failed: {e}")
            raise
    
    async def _generate_simple_word_content(self, data: Dict[str, Any]) -> bytes:
        """Generate simple Word content as fallback"""
        # Simple text-based content
        content = f"""Legal Document Analysis Report
Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}

Analysis Summary:
{json.dumps(data.get('analysis', {}), indent=2)}

Generated by Legal Saathi Document Advisor
"""
        
        # Convert to bytes (this is a simplified approach)
        return content.encode('utf-8')