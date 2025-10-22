"""
Export service for generating PDF and Word comparison reports
"""

import io
import logging
from typing import Dict, Any, Optional
from datetime import datetime
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black, red, orange, green
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting comparison reports to PDF and Word formats"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom styles for PDF generation"""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Title'],
            fontSize=24,
            spaceAfter=30,
            textColor=HexColor('#1e40af'),
            alignment=TA_CENTER
        ))
        
        # Heading style
        self.styles.add(ParagraphStyle(
            name='CustomHeading',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceBefore=20,
            spaceAfter=12,
            textColor=HexColor('#1e40af'),
            leftIndent=0
        ))
        
        # Subheading style
        self.styles.add(ParagraphStyle(
            name='CustomSubheading',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceBefore=15,
            spaceAfter=8,
            textColor=HexColor('#374151'),
            leftIndent=0
        ))
        
        # Risk styles
        self.styles.add(ParagraphStyle(
            name='HighRisk',
            parent=self.styles['Normal'],
            textColor=HexColor('#dc2626'),
            fontSize=10
        ))
        
        self.styles.add(ParagraphStyle(
            name='MediumRisk',
            parent=self.styles['Normal'],
            textColor=HexColor('#ea580c'),
            fontSize=10
        ))
        
        self.styles.add(ParagraphStyle(
            name='LowRisk',
            parent=self.styles['Normal'],
            textColor=HexColor('#16a34a'),
            fontSize=10
        ))
    
    async def export_comparison_report_pdf(self, report_data: Dict[str, Any]) -> bytes:
        """Export comparison report as PDF"""
        try:
            logger.info(f"Generating PDF report for comparison: {report_data.get('comparison_id')}")
            
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Build the PDF content
            story = []
            
            # Title
            story.append(Paragraph("Document Comparison Report", self.styles['CustomTitle']))
            story.append(Spacer(1, 20))
            
            # Executive Summary
            story.append(Paragraph("Executive Summary", self.styles['CustomHeading']))
            
            exec_summary = report_data.get('executive_summary', {})
            story.append(Paragraph(f"<b>Verdict:</b> {exec_summary.get('verdict', 'N/A')}", self.styles['Normal']))
            story.append(Paragraph(f"<b>Recommendation:</b> {exec_summary.get('recommendation', 'N/A')}", self.styles['Normal']))
            story.append(Paragraph(f"<b>Safer Document:</b> {exec_summary.get('safer_document', 'N/A')}", self.styles['Normal']))
            story.append(Paragraph(f"<b>Risk Score Difference:</b> {exec_summary.get('risk_score_difference', 0):.3f}", self.styles['Normal']))
            story.append(Spacer(1, 15))
            
            # Key Insights
            if exec_summary.get('key_insights'):
                story.append(Paragraph("Key Insights", self.styles['CustomSubheading']))
                for insight in exec_summary['key_insights']:
                    story.append(Paragraph(f"• {insight}", self.styles['Normal']))
                story.append(Spacer(1, 15))
            
            # Document Summaries
            story.append(Paragraph("Document Summaries", self.styles['CustomHeading']))
            
            doc_summaries = report_data.get('document_summaries', {})
            if 'document1' in doc_summaries:
                doc1 = doc_summaries['document1']
                story.append(Paragraph("Document 1", self.styles['CustomSubheading']))
                story.append(Paragraph(f"Type: {doc1.get('type', 'N/A')}", self.styles['Normal']))
                if 'overall_risk' in doc1:
                    risk = doc1['overall_risk']
                    story.append(Paragraph(f"Risk Level: {risk.get('level', 'N/A')}", self.styles['Normal']))
                    story.append(Paragraph(f"Risk Score: {risk.get('score', 0):.3f}", self.styles['Normal']))
                story.append(Paragraph(f"Clause Count: {doc1.get('clause_count', 0)}", self.styles['Normal']))
                story.append(Spacer(1, 10))
            
            if 'document2' in doc_summaries:
                doc2 = doc_summaries['document2']
                story.append(Paragraph("Document 2", self.styles['CustomSubheading']))
                story.append(Paragraph(f"Type: {doc2.get('type', 'N/A')}", self.styles['Normal']))
                if 'overall_risk' in doc2:
                    risk = doc2['overall_risk']
                    story.append(Paragraph(f"Risk Level: {risk.get('level', 'N/A')}", self.styles['Normal']))
                    story.append(Paragraph(f"Risk Score: {risk.get('score', 0):.3f}", self.styles['Normal']))
                story.append(Paragraph(f"Clause Count: {doc2.get('clause_count', 0)}", self.styles['Normal']))
                story.append(Spacer(1, 15))
            
            # Visual Differences Summary
            visual_diff = report_data.get('visual_differences', {})
            if visual_diff and 'visual_diff_metadata' in visual_diff:
                story.append(Paragraph("Visual Differences Summary", self.styles['CustomHeading']))
                metadata = visual_diff['visual_diff_metadata']
                story.append(Paragraph(f"Matched Clauses: {metadata.get('matched_clauses', 0)}", self.styles['Normal']))
                story.append(Paragraph(f"Unique to Document 1: {metadata.get('unique_to_doc1', 0)}", self.styles['Normal']))
                story.append(Paragraph(f"Unique to Document 2: {metadata.get('unique_to_doc2', 0)}", self.styles['Normal']))
                story.append(Paragraph(f"Total Changes: {visual_diff.get('total_changes', 0)}", self.styles['Normal']))
                story.append(Paragraph(f"High Impact Changes: {visual_diff.get('high_impact_changes', 0)}", self.styles['Normal']))
                story.append(Spacer(1, 15))
            
            # Key Differences
            key_differences = report_data.get('key_differences', [])
            if key_differences:
                story.append(Paragraph("Key Differences", self.styles['CustomHeading']))
                
                for i, diff in enumerate(key_differences[:10]):  # Limit to top 10
                    story.append(Paragraph(f"{i+1}. {diff.get('description', 'N/A')}", self.styles['CustomSubheading']))
                    story.append(Paragraph(f"Severity: {diff.get('severity', 'N/A').upper()}", self.styles['Normal']))
                    story.append(Paragraph(f"Impact: {diff.get('impact', 'N/A')}", self.styles['Normal']))
                    story.append(Paragraph(f"Recommendation: {diff.get('recommendation', 'N/A')}", self.styles['Normal']))
                    
                    if diff.get('document1_value') or diff.get('document2_value'):
                        story.append(Paragraph(f"Document 1: {diff.get('document1_value', 'Not specified')}", self.styles['Normal']))
                        story.append(Paragraph(f"Document 2: {diff.get('document2_value', 'Not specified')}", self.styles['Normal']))
                    
                    story.append(Spacer(1, 10))
            
            # Clause Analysis (Top 5)
            clause_analysis = report_data.get('clause_analysis', [])
            if clause_analysis:
                story.append(PageBreak())
                story.append(Paragraph("Detailed Clause Analysis", self.styles['CustomHeading']))
                
                for i, clause in enumerate(clause_analysis[:5]):  # Top 5 clauses
                    story.append(Paragraph(f"Clause {i+1}: {clause.get('clause_type', 'N/A').replace('_', ' ').title()}", 
                                         self.styles['CustomSubheading']))
                    story.append(Paragraph(f"Risk Difference: {clause.get('risk_difference', 0):.3f}", self.styles['Normal']))
                    story.append(Paragraph(f"Comparison Notes: {clause.get('comparison_notes', 'N/A')}", self.styles['Normal']))
                    story.append(Paragraph(f"Recommendation: {clause.get('recommendation', 'N/A')}", self.styles['Normal']))
                    
                    if clause.get('semantic_similarity'):
                        story.append(Paragraph(f"Semantic Similarity: {clause['semantic_similarity']:.3f}", self.styles['Normal']))
                    
                    story.append(Spacer(1, 15))
            
            # Processing Metrics
            processing_metrics = report_data.get('processing_metrics', {})
            if processing_metrics:
                story.append(Paragraph("Processing Metrics", self.styles['CustomHeading']))
                story.append(Paragraph(f"Total Processing Time: {processing_metrics.get('total_processing_time', 0):.2f} seconds", self.styles['Normal']))
                story.append(Paragraph(f"Semantic Analysis: {'Enabled' if processing_metrics.get('semantic_analysis_enabled') else 'Disabled'}", self.styles['Normal']))
                story.append(Paragraph(f"Embeddings Used: {processing_metrics.get('embeddings_used', 0)}", self.styles['Normal']))
            
            # Footer
            story.append(Spacer(1, 30))
            story.append(Paragraph(f"Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 
                                 self.styles['Normal']))
            story.append(Paragraph("Generated by LegalSaathi AI Document Analysis Platform", 
                                 self.styles['Normal']))
            
            # Build PDF
            doc.build(story)
            
            # Get the value of the BytesIO buffer and return it
            pdf_data = buffer.getvalue()
            buffer.close()
            
            logger.info("PDF report generated successfully")
            return pdf_data
            
        except Exception as e:
            logger.error(f"Failed to generate PDF report: {e}")
            raise Exception(f"PDF generation failed: {str(e)}")
    
    async def export_comparison_report_word(self, report_data: Dict[str, Any]) -> bytes:
        """Export comparison report as Word document"""
        try:
            logger.info(f"Generating Word report for comparison: {report_data.get('comparison_id')}")
            
            # Create a new Document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Document Comparison Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            
            exec_summary = report_data.get('executive_summary', {})
            
            # Create a table for executive summary
            summary_table = doc.add_table(rows=4, cols=2)
            summary_table.style = 'Table Grid'
            
            summary_data = [
                ('Verdict', exec_summary.get('verdict', 'N/A')),
                ('Safer Document', exec_summary.get('safer_document', 'N/A')),
                ('Risk Score Difference', f"{exec_summary.get('risk_score_difference', 0):.3f}"),
                ('Impact Level', exec_summary.get('impact_level', 'N/A'))
            ]
            
            for i, (label, value) in enumerate(summary_data):
                summary_table.cell(i, 0).text = label
                summary_table.cell(i, 1).text = str(value)
            
            doc.add_paragraph()
            
            # Recommendation
            rec_para = doc.add_paragraph()
            rec_para.add_run('Recommendation: ').bold = True
            rec_para.add_run(exec_summary.get('recommendation', 'N/A'))
            
            # Key Insights
            if exec_summary.get('key_insights'):
                doc.add_heading('Key Insights', level=2)
                for insight in exec_summary['key_insights']:
                    doc.add_paragraph(insight, style='List Bullet')
            
            # Document Summaries
            doc.add_heading('Document Summaries', level=1)
            
            doc_summaries = report_data.get('document_summaries', {})
            
            # Document comparison table
            if 'document1' in doc_summaries and 'document2' in doc_summaries:
                doc1 = doc_summaries['document1']
                doc2 = doc_summaries['document2']
                
                comp_table = doc.add_table(rows=4, cols=3)
                comp_table.style = 'Table Grid'
                
                # Headers
                comp_table.cell(0, 0).text = 'Attribute'
                comp_table.cell(0, 1).text = 'Document 1'
                comp_table.cell(0, 2).text = 'Document 2'
                
                # Data
                comp_table.cell(1, 0).text = 'Type'
                comp_table.cell(1, 1).text = doc1.get('type', 'N/A')
                comp_table.cell(1, 2).text = doc2.get('type', 'N/A')
                
                comp_table.cell(2, 0).text = 'Risk Level'
                comp_table.cell(2, 1).text = doc1.get('overall_risk', {}).get('level', 'N/A')
                comp_table.cell(2, 2).text = doc2.get('overall_risk', {}).get('level', 'N/A')
                
                comp_table.cell(3, 0).text = 'Clause Count'
                comp_table.cell(3, 1).text = str(doc1.get('clause_count', 0))
                comp_table.cell(3, 2).text = str(doc2.get('clause_count', 0))
            
            # Visual Differences Summary
            visual_diff = report_data.get('visual_differences', {})
            if visual_diff and 'visual_diff_metadata' in visual_diff:
                doc.add_heading('Visual Differences Summary', level=1)
                
                metadata = visual_diff['visual_diff_metadata']
                diff_para = doc.add_paragraph()
                diff_para.add_run(f"Matched Clauses: {metadata.get('matched_clauses', 0)}\n")
                diff_para.add_run(f"Unique to Document 1: {metadata.get('unique_to_doc1', 0)}\n")
                diff_para.add_run(f"Unique to Document 2: {metadata.get('unique_to_doc2', 0)}\n")
                diff_para.add_run(f"Total Changes: {visual_diff.get('total_changes', 0)}\n")
                diff_para.add_run(f"High Impact Changes: {visual_diff.get('high_impact_changes', 0)}")
            
            # Key Differences
            key_differences = report_data.get('key_differences', [])
            if key_differences:
                doc.add_heading('Key Differences', level=1)
                
                for i, diff in enumerate(key_differences[:8]):  # Limit to top 8
                    doc.add_heading(f"{i+1}. {diff.get('description', 'N/A')}", level=3)
                    
                    diff_para = doc.add_paragraph()
                    diff_para.add_run('Severity: ').bold = True
                    diff_para.add_run(f"{diff.get('severity', 'N/A').upper()}\n")
                    
                    diff_para.add_run('Impact: ').bold = True
                    diff_para.add_run(f"{diff.get('impact', 'N/A')}\n")
                    
                    diff_para.add_run('Recommendation: ').bold = True
                    diff_para.add_run(diff.get('recommendation', 'N/A'))
                    
                    if diff.get('document1_value') or diff.get('document2_value'):
                        values_para = doc.add_paragraph()
                        values_para.add_run('Document 1: ').bold = True
                        values_para.add_run(f"{diff.get('document1_value', 'Not specified')}\n")
                        values_para.add_run('Document 2: ').bold = True
                        values_para.add_run(diff.get('document2_value', 'Not specified'))
            
            # Clause Analysis Summary
            clause_analysis = report_data.get('clause_analysis', [])
            if clause_analysis:
                doc.add_page_break()
                doc.add_heading('Detailed Clause Analysis', level=1)
                
                for i, clause in enumerate(clause_analysis[:5]):  # Top 5 clauses
                    clause_heading = doc.add_heading(f"Clause {i+1}: {clause.get('clause_type', 'N/A').replace('_', ' ').title()}", level=2)
                    
                    clause_para = doc.add_paragraph()
                    clause_para.add_run('Risk Difference: ').bold = True
                    clause_para.add_run(f"{clause.get('risk_difference', 0):.3f}\n")
                    
                    clause_para.add_run('Comparison Notes: ').bold = True
                    clause_para.add_run(f"{clause.get('comparison_notes', 'N/A')}\n")
                    
                    clause_para.add_run('Recommendation: ').bold = True
                    clause_para.add_run(clause.get('recommendation', 'N/A'))
                    
                    if clause.get('semantic_similarity'):
                        sem_para = doc.add_paragraph()
                        sem_para.add_run('Semantic Similarity: ').bold = True
                        sem_para.add_run(f"{clause['semantic_similarity']:.3f}")
            
            # Processing Metrics
            processing_metrics = report_data.get('processing_metrics', {})
            if processing_metrics:
                doc.add_heading('Processing Metrics', level=1)
                
                metrics_para = doc.add_paragraph()
                metrics_para.add_run(f"Total Processing Time: {processing_metrics.get('total_processing_time', 0):.2f} seconds\n")
                metrics_para.add_run(f"Semantic Analysis: {'Enabled' if processing_metrics.get('semantic_analysis_enabled') else 'Disabled'}\n")
                metrics_para.add_run(f"Embeddings Used: {processing_metrics.get('embeddings_used', 0)}")
            
            # Footer
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_para.add_run(f"Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").italic = True
            footer_para.add_run("Generated by LegalSaathi AI Document Analysis Platform").italic = True
            
            # Save to BytesIO
            buffer = io.BytesIO()
            doc.save(buffer)
            word_data = buffer.getvalue()
            buffer.close()
            
            logger.info("Word report generated successfully")
            return word_data
            
        except Exception as e:
            logger.error(f"Failed to generate Word report: {e}")
            raise Exception(f"Word generation failed: {str(e)}")
    
    def get_supported_formats(self) -> list:
        """Get list of supported export formats"""
        return ['pdf', 'docx', 'word']
    
    async def export_report(self, report_data: Dict[str, Any], format: str = 'pdf') -> bytes:
        """Export report in specified format"""
        format = format.lower()
        
        if format == 'pdf':
            return await self.export_comparison_report_pdf(report_data)
        elif format in ['docx', 'word']:
            return await self.export_comparison_report_word(report_data)
        else:
            raise ValueError(f"Unsupported export format: {format}. Supported formats: {self.get_supported_formats()}")